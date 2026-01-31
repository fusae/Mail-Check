#!/usr/bin/env python3
"""
舆情报告生成器核心模块
Sentiment Analysis Report Generator Core
"""

import pandas as pd
import numpy as np
from datetime import datetime
from typing import Dict, List, Any
from collections import Counter
import re

try:
    import jieba
    JIEBA_AVAILABLE = True
except ImportError:
    JIEBA_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ReportGenerator:
    """舆情报告生成器"""

    def __init__(self):
        self.platform_names = {
            '抖音': '抖音',
            'douyin': '抖音',
            '新浪微博': '微博',
            '微博': '微博',
            'weibo': '微博',
            '微信': '微信',
            'wechat': '微信',
            '新闻网站': '新闻网站',
            'news': '新闻网站',
        }
        self.max_key_events = 10
        self.event_reason_limit = 400
        self.event_content_limit = 800

    def normalize_platform(self, platform: str) -> str:
        """标准化平台名称"""
        platform = platform.lower().strip()
        for key, value in self.platform_names.items():
            if key in platform:
                return value
        return platform

    def generate_report_data(
        self,
        df: pd.DataFrame,
        hospital_name: str,
        report_type: str = "special",
        report_period: str = None
    ) -> Dict[str, Any]:
        """
        生成报告数据

        参数：
        - df: 舆情数据DataFrame
        - hospital_name: 医院名称
        - report_type: 报告类型（special/quarterly/monthly）
        - report_period: 报告周期（如"2026Q1"）
        """
        # 数据预处理
        df = self._preprocess_data(df)

        # 生成各个部分的数据
        report_data = {
            'hospital_name': hospital_name,
            'report_type': report_type,
            'report_period': report_period or self._auto_detect_period(df),
            'generated_time': datetime.now().strftime('%Y年%m月%d日 %H:%M'),
            'summary': self._generate_summary(df),
            'overview': self._generate_overview(df),
            'distribution': self._generate_distribution(df),
            'key_events': self._generate_key_events(df),
            'sentiment': self._generate_sentiment(df),
            'risk_assessment': self._generate_risk_assessment(df),
            'recommendations': self._generate_recommendations(df),
            'appendix': self._generate_appendix(df),
            'raw_dataframe': df  # 保存原始数据供调试使用
        }

        return report_data

    def _preprocess_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """数据预处理"""
        # 复制数据
        df = df.copy()

        # 标准化平台名称
        df['来源_标准'] = df['来源'].apply(self.normalize_platform)

        # 解析时间
        df['创建时间_解析'] = pd.to_datetime(df['创建时间'], errors='coerce')

        # 提取日期和小时
        df['日期'] = df['创建时间_解析'].dt.date
        df['小时'] = df['创建时间_解析'].dt.hour

        # 计算环比（如果有历史数据）
        df['风险分_数值'] = pd.to_numeric(df['风险分'], errors='coerce').fillna(0)

        return df

    def _generate_summary(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成报告摘要"""
        total = len(df)
        high_risk = len(df[df['严重程度'] == 'high'])
        active = len(df[df['状态'] == 'active'])
        avg_risk = df['风险分_数值'].mean()

        # 估算影响人数（根据平台和严重程度）
        estimated_reach = self._estimate_reach(df)

        return {
            'total_events': total,
            'high_risk_events': high_risk,
            'active_events': active,
            'average_risk_score': round(avg_risk, 1),
            'estimated_reach': estimated_reach,
            'peak_time': self._find_peak_time(df),
            'trend': self._analyze_trend(df)
        }

    def _estimate_reach(self, df: pd.DataFrame) -> str:
        """估算影响人数"""
        # 简单估算：每条高风险=10万，中风险=1万，低风险=1000
        high_count = len(df[df['严重程度'] == 'high'])
        medium_count = len(df[df['严重程度'] == 'medium'])
        low_count = len(df[df['严重程度'] == 'low'])

        total = high_count * 100000 + medium_count * 10000 + low_count * 1000

        if total >= 1000000:
            return f"{total // 1000000}M+"
        elif total >= 10000:
            return f"{total // 10000}万+"
        else:
            return f"{total // 1000}千+"

    def _find_peak_time(self, df: pd.DataFrame) -> str:
        """找到传播峰值时间"""
        if len(df) == 0:
            return "未知"

        # 按日期分组
        daily_counts = df.groupby('日期').size()
        if len(daily_counts) == 0:
            return "未知"

        peak_date = daily_counts.idxmax()
        return peak_date.strftime('%Y-%m-%d')

    def _analyze_trend(self, df: pd.DataFrame) -> str:
        """分析趋势"""
        if len(df) < 2:
            return "数据不足"

        # 按日期排序
        df_sorted = df.sort_values('创建时间_解析')

        # 简单判断趋势
        first_half = df_sorted[:len(df_sorted)//2]
        second_half = df_sorted[len(df_sorted)//2:]

        first_count = len(first_half)
        second_count = len(second_half)

        if second_count > first_count * 1.5:
            return "上升"
        elif second_count < first_count * 0.7:
            return "下降"
        else:
            return "平稳"

    def _generate_overview(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成概述数据"""
        total = len(df)

        # 按严重程度统计
        severity_counts = df['严重程度'].value_counts()

        # 按状态统计
        status_counts = df['状态'].value_counts()

        # 按平台统计
        platform_counts = df['来源_标准'].value_counts()

        return {
            'total': total,
            'severity_distribution': {
                'high': int(severity_counts.get('high', 0)),
                'medium': int(severity_counts.get('medium', 0)),
                'low': int(severity_counts.get('low', 0))
            },
            'status_distribution': status_counts.to_dict(),
            'platform_distribution': platform_counts.to_dict(),
            'average_risk_score': round(df['风险分_数值'].mean(), 1)
        }

    def _generate_distribution(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成分布分析"""
        return {
            'time_distribution': self._analyze_time_distribution(df),
            'platform_distribution': self._analyze_platform_distribution(df),
            'type_distribution': self._analyze_type_distribution(df),
            'department_distribution': self._analyze_department_distribution(df)
        }

    def _analyze_time_distribution(self, df: pd.DataFrame) -> Dict[str, Any]:
        """时间分布分析"""
        # 按日期统计
        daily_counts = df.groupby('日期').size().to_dict()

        # 按小时统计
        hourly_counts = df.groupby('小时').size().to_dict()

        # 找出发病时段
        peak_hours = sorted(hourly_counts.items(), key=lambda x: x[1], reverse=True)[:3]

        return {
            'daily_counts': {str(k): v for k, v in daily_counts.items()},
            'hourly_counts': hourly_counts,
            'peak_hours': peak_hours,
            'time_pattern': self._detect_time_pattern(df)
        }

    def _detect_time_pattern(self, df: pd.DataFrame) -> str:
        """检测时间模式"""
        if len(df) == 0:
            return "无数据"

        hour_counts = df.groupby('小时').size()

        # 判断是否夜间集中（22:00-02:00）
        night_hours = [22, 23, 0, 1, 2]
        night_count = sum(hour_counts.get(h, 0) for h in night_hours)

        if night_count > len(df) * 0.3:
            return "夜间集中（22:00-02:00）"
        else:
            return "分散"

    def _analyze_platform_distribution(self, df: pd.DataFrame) -> Dict[str, Any]:
        """平台分布分析"""
        platform_counts = df['来源_标准'].value_counts()
        total = len(df)

        distribution = {}
        for platform, count in platform_counts.items():
            distribution[platform] = {
                'count': int(count),
                'percentage': round(count / total * 100, 1)
            }

        return {
            'distribution': distribution,
            'dominant_platform': platform_counts.idxmax() if len(platform_counts) > 0 else "未知",
            'platform_risk': self._assess_platform_risk(df)
        }

    def _assess_platform_risk(self, df: pd.DataFrame) -> Dict[str, str]:
        """评估平台风险等级"""
        platform_risk = {
            '抖音': '极高',
            '微博': '高',
            '微信': '高',
            '新闻网站': '中高'
        }

        result = {}
        for platform in df['来源_标准'].unique():
            result[platform] = platform_risk.get(platform, '中')

        return result

    def _analyze_type_distribution(self, df: pd.DataFrame) -> Dict[str, Any]:
        """类型分布分析"""
        # 从警示理由和内容中提取类型
        types = []

        for _, row in df.iterrows():
            reason = str(row.get('警示理由', '')).lower()
            content = str(row.get('内容', '')).lower()

            if any(keyword in reason or keyword in content for keyword in ['死亡', '死亡事件', '致死', '治死']):
                types.append('医疗质量-死亡事件')
            elif any(keyword in reason or keyword in content for keyword in ['手术', '治疗', '诊断']):
                types.append('医疗质量-手术/治疗')
            elif any(keyword in reason or keyword in content for keyword in ['服务', '态度', '投诉']):
                types.append('服务质量')
            elif any(keyword in reason or keyword in content for keyword in ['费用', '收费', '钱']):
                types.append('费用相关')
            elif any(keyword in reason or keyword in content for keyword in ['环境', '设施', '停车']):
                types.append('环境设施')
            else:
                types.append('其他')

        type_counts = Counter(types)
        total = len(types)

        distribution = {}
        for type_name, count in type_counts.items():
            distribution[type_name] = {
                'count': count,
                'percentage': round(count / total * 100, 1)
            }

        return {
            'distribution': distribution,
            'main_type': type_counts.most_common(1)[0][0] if type_counts else "未知"
        }

    def _analyze_department_distribution(self, df: pd.DataFrame) -> Dict[str, Any]:
        """科室分布分析"""
        # 从内容中提取科室
        departments = []

        dept_keywords = {
            '心内科': ['心内', '心脏', '心科'],
            '心外科': ['心外'],
            '急诊科': ['急诊'],
            '产科': ['产科', '生产'],
            '儿科': ['儿科', '小儿'],
            '耳鼻喉科': ['耳鼻喉', '耳鼻'],
            '骨科': ['骨科'],
            '外科': ['外科']
        }

        for _, row in df.iterrows():
            content = str(row.get('内容', '')) + str(row.get('警示理由', ''))
            found = False

            for dept, keywords in dept_keywords.items():
                if any(keyword in content for keyword in keywords):
                    departments.append(dept)
                    found = True
                    break

            if not found:
                departments.append('其他/未明确')

        dept_counts = Counter(departments)

        return {
            'department_counts': dict(dept_counts),
            'high_risk_departments': dept_counts.most_common(3)
        }

    def _generate_key_events(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """生成重点事件"""
        # 按风险分排序，取前N个
        top_events = df.nlargest(self.max_key_events, '风险分_数值')

        events = []
        for _, row in top_events.iterrows():
            event = {
                'id': row.get('ID', 'N/A'),
                'title': row.get('标题', '无标题'),
                'platform': row.get('来源_标准', row.get('来源', '未知')),
                'severity': row.get('严重程度', 'unknown'),
                'risk_score': int(row.get('风险分_数值', 0)),
                'status': row.get('状态', 'unknown'),
                'time': str(row.get('创建时间', '未知')),
                'reason': row.get('警示理由', '') or '',
                'content': row.get('内容', '') or '',
                'link': row.get('原文链接', ''),
                'department': self._extract_department(row.get('内容', '')),
                'event_type': self._classify_event(row.get('警示理由', ''), row.get('内容', ''))
            }
            events.append(event)

        return events

    def _extract_department(self, content: str) -> str:
        """从内容中提取科室"""
        dept_keywords = {
            '心内科': ['心内', '心脏', '心科'],
            '心外科': ['心外'],
            '急诊科': ['急诊'],
            '产科': ['产科', '生产'],
            '儿科': ['儿科', '小儿'],
            '耳鼻喉科': ['耳鼻喉', '耳鼻'],
        }

        content = str(content)
        for dept, keywords in dept_keywords.items():
            if any(keyword in content for keyword in keywords):
                return dept

        return "未明确"

    def _classify_event(self, reason: str, content: str) -> str:
        """事件分类"""
        text = (str(reason) + ' ' + str(content)).lower()

        if '死亡' in text or '致死' in text or '治死' in text:
            return '🔴 极高风险 - 患者死亡'
        elif '手术' in text:
            return '🟠 高风险 - 手术相关'
        elif '投诉' in text:
            return '🟡 中风险 - 服务投诉'
        else:
            return '🟢 一般风险'

    def _generate_sentiment(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成情感分析"""
        # 简单情感分析（基于关键词）
        emotions = {'愤怒': 0, '悲伤': 0, '失望': 0, '质疑': 0, '担忧': 0}

        keywords = {
            '愤怒': ['治死', '害死', '不负责任', '垃圾', '无良'],
            '悲伤': ['好好的一个人', '去世', '走了', '难过'],
            '失望': ['失望', '不相信', '怀疑'],
            '质疑': ['质疑', '为什么', '怎么回事'],
            '担忧': ['担心', '害怕', '恐慌']
        }

        for _, row in df.iterrows():
            content = str(row.get('内容', '')).lower()

            for emotion, words in keywords.items():
                if any(word in content for word in words):
                    emotions[emotion] += 1

        total = sum(emotions.values())
        sentiment_distribution = {}
        for emotion, count in emotions.items():
            sentiment_distribution[emotion] = {
                'count': count,
                'percentage': round(count / total * 100, 1) if total > 0 else 0
            }

        # 提取关键词
        top_keywords = self._extract_keywords(df, top_n=20)

        return {
            'sentiment_distribution': sentiment_distribution,
            'dominant_emotion': max(emotions, key=emotions.get) if total > 0 else '未知',
            'sentiment_intensity': '极强' if total > len(df) * 0.5 else '一般',
            'top_keywords': top_keywords
        }

    def _extract_keywords(self, df: pd.DataFrame, top_n: int = 20) -> List[tuple]:
        """提取高频关键词"""
        all_text = ' '.join(df['内容'].fillna('').astype(str).tolist())

        if JIEBA_AVAILABLE:
            words = jieba.cut(all_text)
            # 过滤停用词
            stopwords = {'的', '了', '是', '我', '你', '他', '她', '在', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这'}
            words = [w for w in words if len(w) > 1 and w not in stopwords]
        else:
            # 简单分词（按空格和标点）
            import re
            words = re.findall(r'[\w]{2,}', all_text)

        word_counts = Counter(words)
        return word_counts.most_common(top_n)

    def _generate_risk_assessment(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成风险评估"""
        high_risk = df[df['严重程度'] == 'high']
        medium_risk = df[df['严重程度'] == 'medium']
        low_risk = df[df['严重程度'] == 'low']

        return {
            'current_risks': {
                'critical': {
                    'count': len(high_risk),
                    'events': high_risk.nlargest(3, '风险分_数值')[['标题', '风险分_数值']].to_dict('records') if len(high_risk) > 0 else []
                },
                'high': {
                    'count': len(medium_risk),
                    'events': []
                },
                'medium': {
                    'count': len(low_risk),
                    'events': []
                }
            },
            'risk_level': self._calculate_overall_risk(df),
            'impact_prediction': self._predict_impact(df)
        }

    def _calculate_overall_risk(self, df: pd.DataFrame) -> str:
        """计算总体风险等级"""
        avg_risk = df['风险分_数值'].mean()
        high_ratio = len(df[df['严重程度'] == 'high']) / len(df)

        if avg_risk >= 80 or high_ratio >= 0.5:
            return '🔴 极高危险级别'
        elif avg_risk >= 60 or high_ratio >= 0.3:
            return '🟠 高危险级别'
        elif avg_risk >= 40 or high_ratio >= 0.1:
            return '🟡 中危险级别'
        else:
            return '🟢 低危险级别'

    def _predict_impact(self, df: pd.DataFrame) -> Dict[str, str]:
        """预测影响"""
        return {
            'short_term': '平台持续发酵，可能新增相关内容',
            'medium_term': '可能引发媒体关注和监管介入',
            'long_term': '医院声誉受损，需长期修复',
            'legal_risk': '可能面临法律诉讼，赔偿金额预估50-200万'
        }

    def _generate_recommendations(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成建议"""
        return {
            'immediate_actions': [
                '🚨 立即启动危机公关响应',
                '📢 24小时内发布官方声明',
                '🔍 启动内部调查',
                '🤝 主动与相关方沟通'
            ],
            'short_term_actions': [
                '公布调查进展',
                '处理相关责任人',
                '整改医疗流程'
            ],
            'long_term_actions': [
                '建立危机预警机制',
                '加强医患沟通培训',
                '提升服务质量',
                '定期舆情监测'
            ],
            'monitoring_focus': [
                '抖音平台（最高优先级）',
                '微博',
                '微信',
                '地方论坛'
            ]
        }

    def _generate_appendix(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成附录数据"""
        # 确保必要的列存在
        columns_to_export = ['创建时间', '来源', '严重程度', '风险分', '状态']
        if 'ID' in df.columns:
            columns_to_export.insert(0, 'ID')

        event_list = df[columns_to_export].to_dict('records')

        return {
            'event_list': event_list,
            'statistics': {
                'total': len(df),
                'by_severity': df['严重程度'].value_counts().to_dict(),
                'by_platform': df['来源_标准'].value_counts().to_dict(),
                'by_status': df['状态'].value_counts().to_dict()
            }
        }

    def _auto_detect_period(self, df: pd.DataFrame) -> str:
        """自动检测报告周期"""
        if len(df) == 0:
            return datetime.now().strftime('%YQ%q')

        dates = pd.to_datetime(df['创建时间'], errors='coerce')
        min_date = dates.min()
        max_date = dates.max()

        if pd.isna(min_date) or pd.isna(max_date):
            return datetime.now().strftime('%YQ%q')

        # 判断是季度、月度还是专项报告
        days_diff = (max_date - min_date).days

        if days_diff <= 7:
            return "专项报告"
        elif days_diff <= 31:
            return min_date.strftime('%Y年%m月')
        elif days_diff <= 120:
            quarter = (min_date.month - 1) // 3 + 1
            return f"{min_date.strftime('%Y')}Q{quarter}"
        else:
            return f"{min_date.strftime('%Y')}年度"

    def generate_markdown_report(self, report_data: Dict[str, Any], output_path: str):
        """生成Markdown格式报告"""
        md_content = self._render_markdown_template(report_data)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)

    def _render_markdown_template(self, data: Dict[str, Any]) -> str:
        """渲染Markdown模板"""
        summary = data['summary']
        overview = data['overview']
        dist = data['distribution']
        events = data['key_events']
        sentiment = data['sentiment']
        risk = data['risk_assessment']
        recs = data['recommendations']

        md = f"""# {data['hospital_name']}负面舆情分析报告

**报告周期：** {data['report_period']}
**报告时间：** {data['generated_time']}
**报告类型：** {data['report_type']}

---

## 一、报告概述

### 1.1 舆情总体态势

**{risk['risk_level']}！**

本周期内发现负面舆情{summary['total_events']}条，其中高风险事件{summary['high_risk_events']}条。
估算影响人数：{summary['estimated_reach']}。
传播峰值时间：{summary['peak_time']}。
趋势：{summary['trend']}。

### 1.2 关键数据摘要

| 指标 | 数值 | 说明 |
|------|------|------|
| **负面舆情总数** | {summary['total_events']}条 | 高风险{summary['high_risk_events']}条 |
| **高危事件数量** | {summary['high_risk_events']}起 | 需立即处理 |
| **影响人数估算** | {summary['estimated_reach']} | 平台播放量估算 |
| **传播峰值时间** | {summary['peak_time']} | 集中爆发期 |
| **平均风险分** | {summary['average_risk_score']}分 | 满分100分 |
| **未处理事件** | {summary['active_events']}条 | 状态为active |

---

## 二、舆情分布分析

### 2.1 平台分布

"""

        # 平台分布表格
        for platform, info in dist['platform_distribution']['distribution'].items():
            md += f"- **{platform}**: {info['count']}条 ({info['percentage']}%)\n"

        md += f"""
### 2.2 类型分布

"""

        # 类型分布
        for event_type, info in dist['type_distribution']['distribution'].items():
            md += f"- **{event_type}**: {info['count']}条 ({info['percentage']}%)\n"

        md += f"""
### 2.3 科室分布

高风险科室：
"""

        # 科室分布
        for dept, count in dist['department_distribution']['high_risk_departments']:
            md += f"- **{dept}**: {count}条\n"

        md += f"""

---

## 三、重点负面事件

"""

        # 重点事件
        for i, event in enumerate(events, 1):
            reason = (event.get('reason') or '').strip()
            content = (event.get('content') or '').strip()
            if len(reason) > self.event_reason_limit:
                reason = reason[:self.event_reason_limit] + "..."
            if len(content) > self.event_content_limit:
                content = content[:self.event_content_limit] + "..."
            md += f"""
### {i}. {event['title']}

| 项目 | 详情 |
|------|------|
| **时间** | {event['time']} |
| **平台** | {event['platform']} |
| **类型** | {event['event_type']} |
| **科室** | {event['department']} |
| **风险分** | {event['risk_score']}/100 |
| **状态** | {event['status']} |

**事件概述：**
{reason or '（暂无）'}

**详细内容：**
{content or '（暂无）'}

**原文链接：**
{event.get('link') or '（暂无）'}

"""

        md += f"""

---

## 四、情感分析

### 4.1 情感倾向

"""

        # 情感分布
        for emotion, sentiment_data in sentiment['sentiment_distribution'].items():
            md += f"- **{emotion}**: {sentiment_data['count']}条 ({sentiment_data['percentage']}%)\n"

        md += f"""
**主要情绪：** {sentiment['dominant_emotion']}
**强度：** {sentiment['sentiment_intensity']}

### 4.2 高频关键词

"""

        # 关键词
        for word, count in sentiment['top_keywords'][:15]:
            md += f"- {word} ({count}次)\n"

        md += f"""

---

## 五、风险评估与应对

### 5.1 风险等级

{risk['risk_level']}

### 5.2 立即应对措施（24小时内）

"""

        # 立即措施
        for action in recs['immediate_actions']:
            md += f"{action}\n"

        md += f"""

### 5.3 短期措施（1周内）

"""

        # 短期措施
        for action in recs['short_term_actions']:
            md += f"- {action}\n"

        md += f"""

### 5.4 长期措施（1个月以上）

"""

        # 长期措施
        for action in recs['long_term_actions']:
            md += f"- {action}\n"

        md += f"""

---

## 六、监测重点

"""

        # 监测重点
        for item in recs['monitoring_focus']:
            md += f"- {item}\n"

        md += f"""

---

## 七、附录

### 7.1 事件清单

| 时间 | 平台 | 类型 | 风险分 | 状态 |
|------|------|------|--------|------|
"""

        # 事件清单
        for event in data['appendix']['event_list'][:20]:
            md += f"| {event['创建时间']} | {event['来源']} | {event['严重程度']} | {event['风险分']} | {event['状态']} |\n"

        md += f"""

---

**报告生成时间：** {data['generated_time']}
**报告有效期：** 建议每日更新

---

*本报告基于提供的数据生成，部分信息需核实后使用。*
"""

        return md

    def generate_word_report(self, report_data: Dict[str, Any], output_path: str):
        """生成Word格式报告"""
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装python-docx: pip install python-docx")

        doc = Document()
        summary = report_data['summary']
        dist = report_data['distribution']
        sentiment = report_data['sentiment']
        risk = report_data['risk_assessment']
        recs = report_data['recommendations']
        appendix = report_data['appendix']

        # 标题
        title = doc.add_heading(f"{report_data['hospital_name']}负面舆情分析报告", 0)

        # 报告信息
        info = doc.add_paragraph()
        info.add_run(f"报告周期：{report_data['report_period']}\n")
        info.add_run(f"报告时间：{report_data['generated_time']}\n")
        info.add_run(f"报告类型：{report_data['report_type']}")

        # 概述
        doc.add_heading('一、报告概述', 1)
        doc.add_paragraph(f"风险等级：{risk['risk_level']}")
        doc.add_paragraph(f"负面舆情总数：{summary['total_events']}条")
        doc.add_paragraph(f"高风险事件：{summary['high_risk_events']}条")
        doc.add_paragraph(f"影响人数估算：{summary['estimated_reach']}")
        doc.add_paragraph(f"传播峰值时间：{summary['peak_time']}")
        doc.add_paragraph(f"趋势判断：{summary['trend']}")

        # 重点事件
        doc.add_heading('二、重点事件', 1)
        for event in report_data['key_events']:
            doc.add_heading(event['title'], 2)
            p = doc.add_paragraph()
            p.add_run(f"时间：{event['time']}\n")
            p.add_run(f"平台：{event['platform']}\n")
            p.add_run(f"风险分：{event['risk_score']}/100\n")
            if event.get('reason'):
                reason = event['reason'][:self.event_reason_limit]
                if len(event['reason']) > self.event_reason_limit:
                    reason += "..."
                p.add_run(f"概述：{reason}\n")
            if event.get('content'):
                content = event['content'][:self.event_content_limit]
                if len(event['content']) > self.event_content_limit:
                    content += "..."
                p.add_run(f"内容：{content}\n")
            if event.get('link'):
                p.add_run(f"原文链接：{event['link']}\n")

        # 舆情分布分析
        doc.add_heading('三、舆情分布分析', 1)
        doc.add_heading('3.1 平台分布', 2)
        for platform, info in dist['platform_distribution']['distribution'].items():
            doc.add_paragraph(f"• {platform}: {info['count']}条（{info['percentage']}%）")

        doc.add_heading('3.2 类型分布', 2)
        for event_type, info in dist['type_distribution']['distribution'].items():
            doc.add_paragraph(f"• {event_type}: {info['count']}条（{info['percentage']}%）")

        doc.add_heading('3.3 科室分布', 2)
        for dept, count in dist['department_distribution']['high_risk_departments']:
            doc.add_paragraph(f"• {dept}: {count}条")

        doc.add_heading('3.4 时间分布', 2)
        time_pattern = dist['time_distribution'].get('time_pattern', '未知')
        peak_hours = dist['time_distribution'].get('peak_hours', [])
        if peak_hours:
            peak_desc = "，".join([f"{h}:00({c}条)" for h, c in peak_hours])
        else:
            peak_desc = "暂无"
        doc.add_paragraph(f"时间规律：{time_pattern}")
        doc.add_paragraph(f"高峰时段：{peak_desc}")

        # 情感分析
        doc.add_heading('四、情感分析', 1)
        for emotion, info in sentiment['sentiment_distribution'].items():
            doc.add_paragraph(f"• {emotion}: {info['count']}条（{info['percentage']}%）")
        doc.add_paragraph(f"主要情绪：{sentiment['dominant_emotion']}")
        doc.add_paragraph(f"强度：{sentiment['sentiment_intensity']}")
        if sentiment.get('top_keywords'):
            keywords_text = "、".join([f"{w}({c})" for w, c in sentiment['top_keywords'][:15]])
            doc.add_paragraph(f"高频关键词：{keywords_text}")

        # 风险评估
        doc.add_heading('五、风险评估', 1)
        doc.add_paragraph(f"总体风险等级：{risk['risk_level']}")
        doc.add_paragraph(f"高风险事件数：{risk['current_risks']['critical']['count']}")
        doc.add_paragraph(f"中风险事件数：{risk['current_risks']['high']['count']}")
        doc.add_paragraph(f"低风险事件数：{risk['current_risks']['medium']['count']}")
        impact = risk.get('impact_prediction', {})
        if impact:
            doc.add_paragraph(f"短期影响：{impact.get('short_term', '—')}")
            doc.add_paragraph(f"中期影响：{impact.get('medium_term', '—')}")
            doc.add_paragraph(f"长期影响：{impact.get('long_term', '—')}")
            doc.add_paragraph(f"法律风险：{impact.get('legal_risk', '—')}")

        # 应对措施
        doc.add_heading('六、应对措施', 1)
        doc.add_heading('6.1 立即应对（24小时内）', 2)
        for action in recs['immediate_actions']:
            doc.add_paragraph(action)
        doc.add_heading('6.2 短期措施（1周内）', 2)
        for action in recs['short_term_actions']:
            doc.add_paragraph(f"• {action}")
        doc.add_heading('6.3 长期措施（1个月以上）', 2)
        for action in recs['long_term_actions']:
            doc.add_paragraph(f"• {action}")

        # 监测重点
        doc.add_heading('七、监测重点', 1)
        for item in recs.get('monitoring_focus', []):
            doc.add_paragraph(f"• {item}")

        # 附录数据
        doc.add_heading('八、附录数据', 1)
        event_list = appendix.get('event_list', [])[:20]
        if event_list:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "时间"
            hdr[1].text = "来源"
            hdr[2].text = "严重程度"
            hdr[3].text = "风险分"
            hdr[4].text = "状态"
            for item in event_list:
                row = table.add_row().cells
                row[0].text = str(item.get('创建时间', ''))
                row[1].text = str(item.get('来源', ''))
                row[2].text = str(item.get('严重程度', ''))
                row[3].text = str(item.get('风险分', ''))
                row[4].text = str(item.get('状态', ''))
        else:
            doc.add_paragraph("暂无附录数据")

        # 保存
        doc.save(output_path)


if __name__ == "__main__":
    # 测试
    generator = ReportGenerator()
    print("报告生成器已就绪！")
