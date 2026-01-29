#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
舆情监控 API 服务（Flask）
同域名提供前端数据与 AI 总结/洞察接口
"""

import hashlib
import hmac
import io
import logging
import os
import re
import sqlite3
from collections import defaultdict
from datetime import datetime, timedelta, timezone

import requests
import yaml
from flask import Flask, jsonify, request, send_file

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

logging.basicConfig(level=logging.DEBUG)


app = Flask(__name__)

current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)


def load_config():
    config_path = os.path.join(project_root, 'config', 'config.yaml')
    with open(config_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)


def _get_config_path():
    return os.path.join(project_root, 'config', 'config.yaml')


def _read_config_file():
    config_path = _get_config_path()
    with open(config_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f) or {}


def _write_config_file(cfg):
    config_path = _get_config_path()
    with open(config_path, 'w', encoding='utf-8') as f:
        yaml.safe_dump(cfg, f, allow_unicode=True, sort_keys=False)


config = load_config()
DB_PATH = config.get('runtime', {}).get(
    'database_path',
    os.path.join(project_root, 'data', 'processed_emails.db')
)

ai_config = config.get('ai', {})
feedback_config = config.get('feedback', {})


def _ensure_column(cursor, table_name, column_name, column_def):
    cursor.execute(f"PRAGMA table_info({table_name})")
    columns = {row[1] for row in cursor.fetchall()}
    if column_name not in columns:
        cursor.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_def}")


def init_database():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS negative_sentiments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            sentiment_id TEXT,
            hospital_name TEXT,
            title TEXT,
            source TEXT,
            content TEXT,
            reason TEXT,
            severity TEXT,
            url TEXT,
            status TEXT DEFAULT 'active',
            dismissed_at TEXT,
            processed_at TEXT DEFAULT (datetime('now','localtime'))
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sentiment_feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            sentiment_id TEXT,
            feedback_judgment BOOLEAN,
            feedback_type TEXT,
            feedback_text TEXT,
            user_id TEXT,
            feedback_time TEXT,
            created_at TEXT DEFAULT (datetime('now','localtime'))
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS feedback_rules (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            pattern TEXT,
            rule_type TEXT,
            action TEXT,
            confidence REAL,
            enabled INTEGER DEFAULT 1,
            source_feedback_id INTEGER,
            created_at TEXT DEFAULT (datetime('now','localtime'))
        )
    ''')

    _ensure_column(cursor, 'negative_sentiments', 'content', 'TEXT')
    _ensure_column(cursor, 'negative_sentiments', 'url', 'TEXT')
    _ensure_column(cursor, 'negative_sentiments', 'status', 'TEXT')
    _ensure_column(cursor, 'negative_sentiments', 'dismissed_at', 'TEXT')
    _ensure_column(cursor, 'negative_sentiments', 'insight_text', 'TEXT')
    _ensure_column(cursor, 'negative_sentiments', 'insight_at', 'TEXT')

    cursor.execute('CREATE INDEX IF NOT EXISTS idx_negative_sentiments_processed_at ON negative_sentiments(processed_at)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_negative_sentiments_status ON negative_sentiments(status)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_negative_sentiments_hospital ON negative_sentiments(hospital_name)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_negative_sentiments_sentiment_id ON negative_sentiments(sentiment_id)')

    conn.commit()
    conn.close()


def query_db(sql, params=(), fetchone=False):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute(sql, params)
    rows = cursor.fetchone() if fetchone else cursor.fetchall()
    conn.close()
    return rows


def _now_local_str():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def row_to_opinion(row, include_content=True, preview_len=240):
    content = row["content"] or ""
    if include_content:
        content_out = content
        truncated = False
    else:
        content_out = content[:preview_len]
        truncated = len(content) > preview_len
    return {
        "id": row["sentiment_id"],
        "hospital": row["hospital_name"],
        "title": row["title"],
        "source": row["source"],
        "content": content_out,
        "reason": row["reason"],
        "severity": row["severity"],
        "score": 1.0 if row["severity"] == "high" else 0.6 if row["severity"] == "medium" else 0.3,
        "url": row["url"],
        "status": row["status"] or "active",
        "dismissed_at": row["dismissed_at"],
        "content_truncated": truncated,
        "createdAt": _format_local_time(row["processed_at"]),
    }


def get_sentiment_info(sentiment_id):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT hospital_name, title, source, content, reason, severity, url, status, dismissed_at
        FROM negative_sentiments
        WHERE sentiment_id = ?
    ''', (sentiment_id,))
    result = cursor.fetchone()
    conn.close()

    if result:
        return {
            'hospital_name': result[0],
            'title': result[1],
            'source': result[2],
            'content': result[3],
            'reason': result[4],
            'severity': result[5],
            'url': result[6] if len(result) > 6 else '',
            'status': result[7] if len(result) > 7 else 'active',
            'dismissed_at': result[8] if len(result) > 8 else None
        }
    return None


def verify_signature(sentiment_id, sig):
    secret = feedback_config.get('link_secret', '')
    if not secret:
        return False

    message = f"{sentiment_id}".encode('utf-8')
    expected = hmac.new(secret.encode('utf-8'), message, hashlib.sha256).hexdigest()
    if not hmac.compare_digest(expected, sig or ''):
        return False
    return True


def save_feedback(data):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO sentiment_feedback (
            sentiment_id, feedback_judgment, feedback_type,
            feedback_text, user_id, feedback_time, created_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (
        data['sentiment_id'],
        data['feedback_judgment'],
        data['feedback_type'],
        data['feedback_text'],
        data['user_id'],
        _now_local_str(),
        _now_local_str(),
    ))
    feedback_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return feedback_id


def delete_negative_sentiment(sentiment_id):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        UPDATE negative_sentiments
        SET status = 'dismissed', dismissed_at = ?
        WHERE sentiment_id = ?
    ''', (_now_local_str(), sentiment_id))
    conn.commit()
    conn.close()


def restore_negative_sentiment(sentiment_id):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        UPDATE negative_sentiments
        SET status = 'active', dismissed_at = NULL
        WHERE sentiment_id = ?
    ''', (sentiment_id,))
    conn.commit()
    conn.close()


def get_feedback_list(sentiment_id):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT feedback_time, feedback_type, feedback_text, user_id
        FROM sentiment_feedback
        WHERE sentiment_id = ?
        ORDER BY created_at DESC
        LIMIT 20
    ''', (sentiment_id,))
    rows = cursor.fetchall()
    conn.close()
    return [
        {
            'feedback_time': row[0],
            'feedback_type': row[1],
            'feedback_text': row[2],
            'user_id': row[3]
        }
        for row in rows
    ]


def extract_rule_candidates(text):
    rules = []
    if not text:
        return rules

    explicit_patterns = []
    keyword_match = re.search(r'(关键词|关键字|排除|规则)[:：]\s*(.+)', text)
    if keyword_match:
        explicit_patterns.append(keyword_match.group(2))

    quoted = re.findall(r'[“"《](.+?)[”"》]', text)
    explicit_patterns.extend(quoted)

    candidates = []
    for raw in explicit_patterns:
        parts = re.split(r'[，,、;；\s]+', raw)
        for part in parts:
            term = part.strip()
            if 2 <= len(term) <= 20:
                candidates.append(term)

    for term in candidates:
        rules.append({
            'pattern': term,
            'rule_type': 'keyword',
            'confidence': 0.9
        })

    return rules


def save_feedback_rules(feedback_id, rules, action):
    if not rules:
        return

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    for rule in rules:
        cursor.execute('''
            INSERT INTO feedback_rules (
                pattern, rule_type, action, confidence, enabled, source_feedback_id, created_at
            ) VALUES (?, ?, ?, ?, 1, ?, ?)
        ''', (
            rule.get('pattern'),
            rule.get('rule_type', 'keyword'),
            action,
            rule.get('confidence', 0.5),
            feedback_id,
            _now_local_str(),
        ))
    conn.commit()
    conn.close()


@app.get("/api/opinions")
def list_opinions():
    status = request.args.get("status", "active")
    limit = int(request.args.get("limit", 50))
    offset = int(request.args.get("offset", 0))
    compact = request.args.get("compact", "0").lower() in ("1", "true", "yes")
    preview_len = int(request.args.get("preview", 240))

    logging.debug(
        f"list_opinions called: status={status}, limit={limit}, offset={offset}, compact={compact}"
    )

    rows = query_db(
        """
        SELECT sentiment_id, hospital_name, title, source, content, reason,
               severity, url, status, dismissed_at, processed_at
        FROM negative_sentiments
        WHERE (? = 'all' OR COALESCE(NULLIF(status, ''), 'active') = ?)
        ORDER BY processed_at DESC
        LIMIT ? OFFSET ?
        """,
        (status, status, limit, offset),
    )
    logging.debug(f"Query returned {len(rows)} rows")
    return jsonify([row_to_opinion(r, include_content=not compact, preview_len=preview_len) for r in rows])


@app.get("/api/stats")
def get_stats():
    range_key = request.args.get("range", "7d")
    start_date = request.args.get("start_date", "")
    end_date = request.args.get("end_date", "")

    now = datetime.now()
    if start_date or end_date:
        start_dt = _parse_db_datetime(start_date) if start_date else None
        end_dt = _parse_db_datetime(end_date) if end_date else None
    else:
        if range_key == "24h":
            start_dt = now - timedelta(hours=24)
        elif range_key == "30d":
            start_dt = now - timedelta(days=30)
        else:
            start_dt = now - timedelta(days=7)
        end_dt = now

    time_filters = []
    params = []
    if start_dt:
        time_filters.append("processed_at >= ?")
        params.append(start_dt.strftime("%Y-%m-%d %H:%M:%S"))
    if end_dt:
        time_filters.append("processed_at <= ?")
        params.append(end_dt.strftime("%Y-%m-%d %H:%M:%S"))
    time_clause = " AND ".join(time_filters)
    active_clause = " AND ".join([f for f in [time_clause, "COALESCE(NULLIF(status,''),'active') != 'dismissed'"] if f])
    time_where = f"WHERE {time_clause}" if time_clause else ""
    active_where = f"WHERE {active_clause}" if active_clause else ""

    proc_cond = ""
    dis_cond = ""
    stats_params = {}
    if start_dt:
        proc_cond += " AND processed_at >= :start_proc"
        dis_cond += " AND dismissed_at >= :start_dis"
        stats_params["start_proc"] = start_dt.strftime("%Y-%m-%d %H:%M:%S")
        stats_params["start_dis"] = start_dt.strftime("%Y-%m-%d %H:%M:%S")
    if end_dt:
        proc_cond += " AND processed_at <= :end_proc"
        dis_cond += " AND dismissed_at <= :end_dis"
        stats_params["end_proc"] = end_dt.strftime("%Y-%m-%d %H:%M:%S")
        stats_params["end_dis"] = end_dt.strftime("%Y-%m-%d %H:%M:%S")

    rows = query_db(
        f"""
        SELECT
            SUM(CASE WHEN COALESCE(NULLIF(status,''),'active') != 'dismissed'{proc_cond} THEN 1 ELSE 0 END) AS active_total,
            SUM(CASE WHEN COALESCE(NULLIF(status,''),'active') = 'dismissed'{dis_cond} THEN 1 ELSE 0 END) AS dismissed_total,
            SUM(CASE WHEN severity = 'high' AND COALESCE(NULLIF(status,''),'active') != 'dismissed'{proc_cond} THEN 1 ELSE 0 END) AS high_total,
            SUM(CASE WHEN severity = 'medium' AND COALESCE(NULLIF(status,''),'active') != 'dismissed'{proc_cond} THEN 1 ELSE 0 END) AS medium_total,
            SUM(CASE WHEN severity = 'low' AND COALESCE(NULLIF(status,''),'active') != 'dismissed'{proc_cond} THEN 1 ELSE 0 END) AS low_total,
            SUM(CASE WHEN COALESCE(NULLIF(status,''),'active') != 'dismissed'{proc_cond}
                THEN CASE
                    WHEN severity = 'high' THEN 0.92
                    WHEN severity = 'medium' THEN 0.6
                    ELSE 0.35
                END
                ELSE 0 END) AS total_score,
            SUM(CASE WHEN COALESCE(NULLIF(status,''),'active') != 'dismissed'{proc_cond} THEN 1 ELSE 0 END) AS score_count
        FROM negative_sentiments
        """,
        stats_params,
        fetchone=True,
    )
    sources_rows = query_db(
        f"""
        SELECT
            COALESCE(NULLIF(source,''),'未知') AS source,
            COUNT(*) AS count
        FROM negative_sentiments
        {active_where}
        GROUP BY COALESCE(NULLIF(source,''),'未知')
        ORDER BY count DESC
        LIMIT 10
        """,
        tuple(params),
    )
    hospital_list_rows = query_db(
        f"""
        SELECT DISTINCT COALESCE(NULLIF(hospital_name,''),'未知') AS hospital
        FROM negative_sentiments
        {active_where}
        ORDER BY hospital
        """,
        tuple(params),
    )
    hospital_rows = query_db(
        f"""
        SELECT
            COALESCE(NULLIF(hospital_name,''),'未知') AS hospital,
            SUM(CASE WHEN severity = 'high' THEN 1 ELSE 0 END) AS high,
            SUM(CASE WHEN severity = 'medium' THEN 1 ELSE 0 END) AS medium,
            SUM(CASE WHEN severity = 'low' THEN 1 ELSE 0 END) AS low,
            COUNT(*) AS total
        FROM negative_sentiments
        {active_where}
        GROUP BY COALESCE(NULLIF(hospital_name,''),'未知')
        ORDER BY total DESC
        LIMIT 10
        """,
        tuple(params),
    )
    score_count = rows["score_count"] or 0
    avg_score = round((rows["total_score"] or 0) / score_count * 100, 1) if score_count else 0
    return jsonify({
        "active_total": rows["active_total"] or 0,
        "dismissed_total": rows["dismissed_total"] or 0,
        "high_total": rows["high_total"] or 0,
        "avg_score": avg_score,
        "severity": {
            "high": rows["high_total"] or 0,
            "medium": rows["medium_total"] or 0,
            "low": rows["low_total"] or 0,
        },
        "sources": [
            {"source": r["source"], "count": r["count"] or 0}
            for r in sources_rows
        ],
        "hospital_list": [r["hospital"] for r in hospital_list_rows],
        "hospitals": [
            {
                "hospital": r["hospital"],
                "high": r["high"] or 0,
                "medium": r["medium"] or 0,
                "low": r["low"] or 0,
                "total": r["total"] or 0,
            }
            for r in hospital_rows
        ],
    })


@app.get("/api/stats/trend")
def get_trend():
    range_key = request.args.get("range", "7d")
    now = datetime.now()
    if range_key == "24h":
        start_dt = now - timedelta(hours=24)
        bucket_fmt = "%H:00"
    elif range_key == "30d":
        start_dt = now - timedelta(days=30)
        bucket_fmt = "%m-%d"
    else:
        start_dt = now - timedelta(days=7)
        bucket_fmt = "%m-%d"

    rows = query_db(
        """
        SELECT processed_at, severity
        FROM negative_sentiments
        WHERE COALESCE(NULLIF(status,''),'active') != 'dismissed'
          AND processed_at >= ?
          AND processed_at <= ?
        ORDER BY processed_at ASC
        """,
        (start_dt.strftime("%Y-%m-%d %H:%M:%S"), now.strftime("%Y-%m-%d %H:%M:%S")),
    )

    buckets = defaultdict(lambda: {"count": 0, "score": 0.0})
    cursor = start_dt
    step = timedelta(hours=1) if range_key == "24h" else timedelta(days=1)
    while cursor <= now:
        buckets[cursor.strftime(bucket_fmt)]
        cursor += step

    for row in rows:
        ts = _parse_db_datetime(row["processed_at"])
        if not ts:
            continue
        label = ts.strftime(bucket_fmt)
        buckets[label]["count"] += 1
        buckets[label]["score"] += _severity_score(row["severity"] or "low")

    data = []
    for label, item in sorted(buckets.items(), key=lambda x: x[0]):
        avg_score = round((item["score"] / item["count"]) * 100) if item["count"] else 0
        data.append({
            "label": label,
            "count": item["count"],
            "avgScore": avg_score,
        })

    return jsonify({"range": range_key, "data": data})


@app.get("/api/opinions/<sentiment_id>")
def get_opinion(sentiment_id):
    row = query_db(
        """
        SELECT sentiment_id, hospital_name, title, source, content, reason,
               severity, url, status, dismissed_at, processed_at
        FROM negative_sentiments
        WHERE sentiment_id = ?
        """,
        (sentiment_id,),
        fetchone=True,
    )
    if not row:
        return jsonify({"error": "not_found"}), 404
    return jsonify(row_to_opinion(row))


@app.get("/api/search")
def search_opinions():
    query = request.args.get("query", "").strip()
    limit = int(request.args.get("limit", 50))
    offset = int(request.args.get("offset", 0))
    compact = request.args.get("compact", "0").lower() in ("1", "true", "yes")
    preview_len = int(request.args.get("preview", 240))

    if not query:
        return jsonify([])

    like = f"%{query}%"
    rows = query_db(
        """
        SELECT sentiment_id, hospital_name, title, source, content, reason,
               severity, url, status, dismissed_at, processed_at
        FROM negative_sentiments
        WHERE hospital_name LIKE ? OR title LIKE ? OR content LIKE ? OR source LIKE ?
        ORDER BY processed_at DESC
        LIMIT ? OFFSET ?
        """,
        (like, like, like, like, limit, offset),
    )
    return jsonify([row_to_opinion(r, include_content=not compact, preview_len=preview_len) for r in rows])


def call_ai(prompt):
    if not ai_config:
        return "AI 未配置"

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {ai_config.get('api_key', '')}"
    }
    data = {
        "model": ai_config.get("model"),
        "messages": [
            {"role": "system", "content": "你是专业的舆情分析助手。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": ai_config.get("temperature", 0.3),
        "max_tokens": ai_config.get("max_tokens", 800)
    }

    api_url = ai_config.get("api_url")
    timeout = ai_config.get("timeout", 60)
    max_retries = ai_config.get("max_retries", 2)
    last_exc = None

    for attempt in range(1, max_retries + 2):
        try:
            resp = requests.post(api_url, headers=headers, json=data, timeout=timeout)
            resp.raise_for_status()
            result = resp.json()
            return result["choices"][0]["message"]["content"]
        except Exception as exc:
            last_exc = exc
            logging.warning(f"AI调用失败（第{attempt}次）: {exc}")
            continue

    raise last_exc


def _parse_db_datetime(value):
    if not value:
        return None
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(value, fmt)
        except ValueError:
            continue
    try:
        return datetime.fromisoformat(value)
    except ValueError:
        return None


def _format_local_time(value):
    dt = _parse_db_datetime(value)
    return dt.strftime("%Y-%m-%d %H:%M:%S") if dt else value


def _severity_score(severity):
    if severity == "high":
        return 0.92
    if severity == "medium":
        return 0.6
    return 0.35


def _query_report_data(hospital, start_date, end_date, include_dismissed=False):
    filters = []
    params = []

    if not include_dismissed:
        filters.append("COALESCE(NULLIF(status, ''), 'active') != 'dismissed'")

    if hospital and hospital != "all":
        filters.append("hospital_name = ?")
        params.append(hospital)

    if start_date:
        filters.append("processed_at >= ?")
        params.append(start_date)
    if end_date:
        filters.append("processed_at <= ?")
        params.append(end_date)

    where_clause = " AND ".join(filters)
    if where_clause:
        where_clause = "WHERE " + where_clause

    rows = query_db(
        f"""
        SELECT sentiment_id, hospital_name, title, source, content, reason,
               severity, url, status, dismissed_at, processed_at
        FROM negative_sentiments
        {where_clause}
        ORDER BY processed_at DESC
        """,
        tuple(params),
    )
    return [row_to_opinion(r) for r in rows]


def _group_by_date(opinions, start_dt, end_dt):
    grouped = defaultdict(int)
    for item in opinions:
        ts = _parse_db_datetime(item.get("createdAt"))
        if not ts:
            continue
        key = ts.strftime("%m-%d")
        grouped[key] += 1

    if start_dt and end_dt:
        cursor = start_dt
        while cursor <= end_dt:
            key = cursor.strftime("%m-%d")
            grouped.setdefault(key, 0)
            cursor += timedelta(days=1)

    return sorted(grouped.items(), key=lambda x: x[0])


def _build_report_charts(opinions, hospital):
    charts = {}

    start_dt = min(
        [d for d in (_parse_db_datetime(o.get("createdAt")) for o in opinions) if d],
        default=None,
    )
    end_dt = max(
        [d for d in (_parse_db_datetime(o.get("createdAt")) for o in opinions) if d],
        default=None,
    )

    # 趋势图
    trend = _group_by_date(opinions, start_dt, end_dt)
    if trend:
        labels, values = zip(*trend)
    else:
        labels, values = [], []
    fig, ax = plt.subplots(figsize=(6, 2.6))
    ax.plot(labels, values, color="#6366f1", linewidth=2)
    ax.fill_between(range(len(values)), values, color="#6366f1", alpha=0.1)
    ax.set_title("舆情走势")
    ax.set_ylabel("数量")
    ax.tick_params(axis="x", rotation=45)
    ax.grid(alpha=0.2)
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    charts["trend"] = buf

    # 来源分布
    source_count = defaultdict(int)
    for item in opinions:
        source_count[item.get("source") or "未知"] += 1
    sources = list(source_count.keys())
    values = list(source_count.values())
    fig, ax = plt.subplots(figsize=(6, 2.6))
    ax.barh(sources, values, color="#10b981")
    ax.set_title("来源分布")
    ax.set_xlabel("数量")
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    charts["source"] = buf

    # 严重程度分布
    severity_count = {
        "低危": len([o for o in opinions if o.get("severity") == "low"]),
        "中危": len([o for o in opinions if o.get("severity") == "medium"]),
        "高危": len([o for o in opinions if o.get("severity") == "high"]),
    }
    fig, ax = plt.subplots(figsize=(6, 2.6))
    ax.bar(severity_count.keys(), severity_count.values(), color=["#10b981", "#f97316", "#ef4444"])
    ax.set_title("严重程度分布")
    ax.set_ylabel("数量")
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    charts["severity"] = buf

    # 医院舆情对比（全院汇总时）
    hospital_count = defaultdict(lambda: {"low": 0, "medium": 0, "high": 0})
    for item in opinions:
        name = item.get("hospital") or "未知"
        hospital_count[name][item.get("severity") or "low"] += 1

    hospitals_sorted = sorted(hospital_count.items(), key=lambda x: sum(x[1].values()), reverse=True)[:8]
    names = [h[0] for h in hospitals_sorted]
    low_vals = [h[1]["low"] for h in hospitals_sorted]
    med_vals = [h[1]["medium"] for h in hospitals_sorted]
    high_vals = [h[1]["high"] for h in hospitals_sorted]

    fig, ax = plt.subplots(figsize=(6, 2.8))
    ax.bar(names, low_vals, color="#10b981", label="低危")
    ax.bar(names, med_vals, bottom=low_vals, color="#f97316", label="中危")
    bottom_high = [low_vals[i] + med_vals[i] for i in range(len(names))]
    ax.bar(names, high_vals, bottom=bottom_high, color="#ef4444", label="高危")
    ax.set_title("医院舆情对比")
    ax.tick_params(axis="x", rotation=30)
    ax.legend()
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    charts["hospital"] = buf

    return charts


def _build_ai_advice(opinions, hospital, start_date, end_date):
    if not opinions:
        return "暂无舆情记录，无需出具处置建议。"

    lines = []
    for idx, op in enumerate(opinions[:10], 1):
        lines.append(
            f"{idx}. 医院:{op.get('hospital')} 标题:{op.get('title')} "
            f"来源:{op.get('source')} 严重程度:{op.get('severity')} "
            f"内容:{(op.get('content') or '')[:120]}"
        )

    prompt = (
        "请基于以下舆情列表给出一段“处置建议”，不需要提及AI或模型来源，"
        "控制在300字以内，可使用条目。\n"
        f"医院范围:{hospital or '全院汇总'}\n"
        f"时间范围:{start_date or '全部'} 至 {end_date or '全部'}\n"
        "舆情列表：\n" + "\n".join(lines)
    )

    try:
        text = call_ai(prompt).strip()
        # 避免在报告中出现“AI/模型”字样
        text = re.sub(r"(AI|大模型|模型|人工智能)\s*", "", text, flags=re.IGNORECASE)
        return text.strip() or "建议：优先处理高危舆情，建立跨部门响应机制，及时澄清事实并跟进患者沟通。"
    except Exception as exc:
        logging.warning(f"AI建议生成失败: {exc}")
        return "建议：优先处理高危舆情，建立跨部门响应机制，及时澄清事实并跟进患者沟通。"


def _get_reportlab_font():
    font_candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/SimHei.ttf",
        "/Library/Fonts/Songti.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
    ]
    for path in font_candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("CJKFont", path))
                return "CJKFont"
            except Exception:
                continue
    return "Helvetica"


def _build_docx_report(opinions, charts, hospital, start_date, end_date, advice_text):
    doc = Document()
    doc.add_heading("舆情监控报告", level=1)
    doc.add_paragraph(f"医院范围：{hospital or '全院汇总'}")
    doc.add_paragraph(f"统计时间：{start_date or '全部'} 至 {end_date or '全部'}")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("一、执行摘要", level=2)
    total = len(opinions)
    high = len([o for o in opinions if o.get("severity") == "high"])
    medium = len([o for o in opinions if o.get("severity") == "medium"])
    low = len([o for o in opinions if o.get("severity") == "low"])
    avg_score = round(
        (sum(_severity_score(o.get("severity")) for o in opinions) / total) * 100, 1
    ) if total else 0
    doc.add_paragraph(f"舆情总量：{total} 条")
    doc.add_paragraph(f"高危：{high} 条 | 中危：{medium} 条 | 低危：{low} 条")
    doc.add_paragraph(f"平均风险指数：{avg_score}")

    doc.add_heading("二、趋势与分布", level=2)
    if charts.get("trend"):
        doc.add_paragraph("舆情走势")
        doc.add_picture(charts["trend"], width=Inches(5.8))
    if charts.get("severity"):
        doc.add_paragraph("严重程度分布")
        doc.add_picture(charts["severity"], width=Inches(5.2))
    if charts.get("source"):
        doc.add_paragraph("来源分布")
        doc.add_picture(charts["source"], width=Inches(5.2))
    if charts.get("hospital"):
        doc.add_paragraph("医院舆情对比")
        doc.add_picture(charts["hospital"], width=Inches(5.8))

    doc.add_heading("三、重点舆情摘要", level=2)
    top_items = sorted(
        opinions,
        key=lambda o: (
            2 if o.get("severity") == "high" else 1 if o.get("severity") == "medium" else 0,
            o.get("createdAt") or "",
        ),
        reverse=True,
    )[:10]
    for idx, item in enumerate(top_items, 1):
        doc.add_paragraph(
            f"{idx}. {item.get('title')}（{item.get('hospital')} / {item.get('source')}）"
        )
        doc.add_paragraph(f"严重程度：{item.get('severity')}  时间：{item.get('createdAt')}")
        doc.add_paragraph(f"警示理由：{item.get('reason')}")
        doc.add_paragraph(f"内容：{item.get('content')}")
        if item.get("url"):
            doc.add_paragraph(f"原文链接：{item.get('url')}")

    doc.add_heading("四、处置建议", level=2)
    doc.add_paragraph(advice_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _build_pdf_report(opinions, charts, hospital, start_date, end_date, advice_text):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=1.5 * cm, leftMargin=1.5 * cm, topMargin=1.5 * cm, bottomMargin=1.5 * cm)

    font_name = _get_reportlab_font()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CJK", fontName=font_name, fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="CJKTitle", fontName=font_name, fontSize=16, leading=20, spaceAfter=12))
    styles.add(ParagraphStyle(name="CJKHeading", fontName=font_name, fontSize=12, leading=16, spaceBefore=10, spaceAfter=6))

    story = []
    story.append(Paragraph("舆情监控报告", styles["CJKTitle"]))
    story.append(Paragraph(f"医院范围：{hospital or '全院汇总'}", styles["CJK"]))
    story.append(Paragraph(f"统计时间：{start_date or '全部'} 至 {end_date or '全部'}", styles["CJK"]))
    story.append(Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["CJK"]))
    story.append(Spacer(1, 8))

    total = len(opinions)
    high = len([o for o in opinions if o.get("severity") == "high"])
    medium = len([o for o in opinions if o.get("severity") == "medium"])
    low = len([o for o in opinions if o.get("severity") == "low"])
    avg_score = round(
        (sum(_severity_score(o.get("severity")) for o in opinions) / total) * 100, 1
    ) if total else 0

    story.append(Paragraph("一、执行摘要", styles["CJKHeading"]))
    summary_table = Table(
        [["舆情总量", total], ["高危", high], ["中危", medium], ["低危", low], ["平均风险指数", avg_score]],
        colWidths=[4 * cm, 4 * cm],
    )
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5f5")),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
            ]
        )
    )
    story.append(summary_table)
    story.append(Spacer(1, 10))

    story.append(Paragraph("二、趋势与分布", styles["CJKHeading"]))
    for key, label in [("trend", "舆情走势"), ("severity", "严重程度分布"), ("source", "来源分布"), ("hospital", "医院舆情对比")]:
        if charts.get(key):
            story.append(Paragraph(label, styles["CJK"]))
            img = RLImage(charts[key], width=16 * cm, height=6 * cm)
            story.append(img)
            story.append(Spacer(1, 6))

    story.append(Paragraph("三、重点舆情摘要", styles["CJKHeading"]))
    top_items = sorted(
        opinions,
        key=lambda o: (
            2 if o.get("severity") == "high" else 1 if o.get("severity") == "medium" else 0,
            o.get("createdAt") or "",
        ),
        reverse=True,
    )[:10]
    for idx, item in enumerate(top_items, 1):
        story.append(Paragraph(f"{idx}. {item.get('title')}（{item.get('hospital')} / {item.get('source')}）", styles["CJK"]))
        story.append(Paragraph(f"严重程度：{item.get('severity')}  时间：{item.get('createdAt')}", styles["CJK"]))
        story.append(Paragraph(f"警示理由：{item.get('reason')}", styles["CJK"]))
        story.append(Paragraph(f"内容：{item.get('content')}", styles["CJK"]))
        if item.get("url"):
            story.append(Paragraph(f"原文链接：{item.get('url')}", styles["CJK"]))
        story.append(Spacer(1, 6))

    story.append(Paragraph("四、处置建议", styles["CJKHeading"]))
    story.append(Paragraph(advice_text, styles["CJK"]))

    doc.build(story)
    buffer.seek(0)
    return buffer


@app.post("/api/ai/summary")
def ai_summary():
    payload = request.get_json(force=True) or {}
    opinions = payload.get("opinions") or []

    if not opinions:
        return jsonify({"text": "暂无负面舆情可总结。"})

    lines = []
    for idx, op in enumerate(opinions, 1):
        lines.append(f"{idx}. 医院:{op.get('hospital')} 标题:{op.get('title')} 内容:{(op.get('content') or '')[:200]}")

    prompt = (
        "请基于以下舆情列表生成一段“现状综述”和“公关建议”。\n"
        "输出格式：\n"
        "现状综述：...\n"
        "公关建议：...\n\n"
        "舆情列表：\n" + "\n".join(lines)
    )

    text = call_ai(prompt)
    return jsonify({"text": text, "generated_at": _now_local_str()})


@app.post("/api/ai/insight")
def ai_insight():
    payload = request.get_json(force=True) or {}
    opinion = payload.get("opinion") or {}

    if not opinion:
        return jsonify({"text": "未提供舆情内容。"}), 400

    sentiment_id = opinion.get("id")
    if sentiment_id:
        cached = query_db(
            "SELECT insight_text, insight_at FROM negative_sentiments WHERE sentiment_id = ?",
            (sentiment_id,),
            fetchone=True,
        )
        if cached and cached["insight_text"]:
            return jsonify({
                "text": cached["insight_text"],
                "generated_at": cached["insight_at"] or _now_local_str(),
                "cached": True,
            })

    prompt = (
        "请对以下单条舆情进行传播风险点分析，并给出简要建议（100字以内）。\n"
        f"医院:{opinion.get('hospital')}\n"
        f"来源:{opinion.get('source')}\n"
        f"标题:{opinion.get('title')}\n"
        f"内容:{opinion.get('content')}\n"
    )
    text = call_ai(prompt)
    generated_at = _now_local_str()
    if sentiment_id:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE negative_sentiments SET insight_text = ?, insight_at = ? WHERE sentiment_id = ?",
            (text, generated_at, sentiment_id),
        )
        conn.commit()
        conn.close()
    return jsonify({"text": text, "generated_at": generated_at, "cached": False})


@app.get("/api/notification/suppress_keywords")
def get_suppress_keywords():
    cfg = _read_config_file()
    notification = cfg.get("notification", {}) or {}
    wechat_cfg = notification.get("wechat_work", {}) or {}
    keywords = wechat_cfg.get("suppress_keywords") or notification.get("suppress_keywords") or []
    if isinstance(keywords, str):
        keywords = [keywords]
    keywords = [str(k).strip() for k in keywords if str(k).strip()]
    return jsonify({"keywords": keywords})


@app.post("/api/notification/suppress_keywords")
def update_suppress_keywords():
    payload = request.get_json(force=True) or {}
    keywords = payload.get("keywords") or []
    if isinstance(keywords, str):
        keywords = [keywords]
    cleaned = []
    for k in keywords:
        text = str(k).strip()
        if text and text not in cleaned:
            cleaned.append(text)

    cfg = _read_config_file()
    notification = cfg.setdefault("notification", {})
    wechat_cfg = notification.setdefault("wechat_work", {})
    wechat_cfg["suppress_keywords"] = cleaned

    _write_config_file(cfg)
    return jsonify({"success": True, "keywords": cleaned})


@app.post("/api/reports")
def export_report():
    payload = request.get_json(force=True) or {}
    hospital = payload.get("hospital") or "all"
    start_date = payload.get("start_date") or ""
    end_date = payload.get("end_date") or ""
    report_format = (payload.get("format") or "pdf").lower()
    include_dismissed = bool(payload.get("include_dismissed", False))

    if start_date and len(start_date) == 10:
        start_date = f"{start_date} 00:00:00"
    if end_date and len(end_date) == 10:
        end_date = f"{end_date} 23:59:59"

    opinions = _query_report_data(hospital, start_date, end_date, include_dismissed)
    charts = _build_report_charts(opinions, hospital)
    advice_text = _build_ai_advice(opinions, hospital, start_date, end_date)

    date_label = datetime.now().strftime("%Y%m%d")
    hospital_label = hospital if hospital != "all" else "全院汇总"

    if report_format in ("doc", "docx", "word"):
        buffer = _build_docx_report(opinions, charts, hospital_label, start_date, end_date, advice_text)
        filename = f"{hospital_label}_舆情报告_{date_label}.docx"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if report_format == "pdf":
        buffer = _build_pdf_report(opinions, charts, hospital_label, start_date, end_date, advice_text)
        filename = f"{hospital_label}_舆情报告_{date_label}.pdf"
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype="application/pdf")

    return jsonify({"error": "unsupported_format"}), 400


@app.get('/feedback')
def feedback_form():
    sentiment_id = request.args.get('sentiment_id', '')
    sig = request.args.get('sig', '')

    if not verify_signature(sentiment_id, sig):
        return "Invalid or expired link.", 403

    sentiment_info = get_sentiment_info(sentiment_id)

    if sentiment_info:
        hospital_name = sentiment_info['hospital_name'] or ''
        title = sentiment_info['title'] or ''
        source = sentiment_info['source'] or ''
        content = sentiment_info['content'] or ''
        reason = sentiment_info['reason'] or ''
        severity = sentiment_info['severity'] or 'medium'
        url = sentiment_info.get('url', '')
        status = sentiment_info.get('status', 'active')
        dismissed_at = sentiment_info.get('dismissed_at')
        feedback_list = get_feedback_list(sentiment_id)

        severity_colors = {
            'high': '#ff4d4f',
            'medium': '#faad14',
            'low': '#52c41a'
        }
        severity_label = {
            'high': '高',
            'medium': '中',
            'low': '低'
        }
        severity_color = severity_colors.get(severity, '#faad14')
        severity_text = severity_label.get(severity, '中')

        feedback_items = ""
        if feedback_list:
            for item in feedback_list:
                feedback_items += f"""
                <div class="feedback-item">
                  <div class="feedback-meta">{item['feedback_time']} | {item['feedback_type']} | {item['user_id']}</div>
                  <div class="feedback-text">{item['feedback_text']}</div>
                </div>
                """
        else:
            feedback_items = "<div class=\"feedback-empty\">暂无反馈</div>"

        status_line = ""
        restore_button = ""
        if status == 'dismissed':
            dismissed_label = dismissed_at or '未知时间'
            status_line = f"""
          <div class="info-row">
            <span class="label">状态：</span>
            <span class="value" style="color: #52c41a; font-weight: bold;">已标记为误报（{dismissed_label}）</span>
          </div>
            """
            restore_button = """
      <button class="btn btn-restore" type="submit" name="action" value="restore">↺ 恢复为负面</button>
            """

        info_section = f"""
        <div class="info-section">
          <h4>舆情详情</h4>
          <div class="info-row">
            <span class="label">舆情ID：</span>
            <span class="value">{sentiment_id}</span>
          </div>
          {status_line}
          <div class="info-row">
            <span class="label">医院：</span>
            <span class="value">{hospital_name}</span>
          </div>
          <div class="info-row">
            <span class="label">来源：</span>
            <span class="value">{source}</span>
          </div>
          <div class="info-row">
            <span class="label">严重程度：</span>
            <span class="value" style="color: {severity_color}; font-weight: bold;">{severity_text}</span>
          </div>
          <div class="info-row">
            <span class="label">标题：</span>
            <span class="value">{title}</span>
          </div>
          <div class="info-row">
            <span class="label">原文链接：</span>
            <span class="value"><a href="{url}" target="_blank" style="color: #1890ff;">{url or '无'}</a></span>
          </div>
          <div class="info-row">
            <span class="label">内容：</span>
            <span class="value">{content}</span>
          </div>
          <div class="info-row">
            <span class="label">AI判断：</span>
            <span class="value">{reason}</span>
          </div>
        </div>
        <div class="feedback-section">
          <h4>用户反馈</h4>
          {feedback_items}
        </div>
        """
    else:
        info_section = f"""
        <div class="info-section" style="background: #fffbe6; border-color: #ffe58f;">
          <h4>舆情详情</h4>
          <p style="color: #faad14;">⚠️ 测试数据（未存入数据库）</p>
          <div class="info-row">
            <span class="label">舆情ID：</span>
            <span class="value">{sentiment_id}</span>
          </div>
          <div class="info-row">
            <span class="label">说明：</span>
            <span class="value">此舆情数据为测试生成，未存入数据库。如需测试完整功能，请通过实际监控产生舆情数据。</span>
          </div>
        </div>
        """

    return f"""<!DOCTYPE html>
<html lang="zh">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>舆情反馈</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 24px; background: #f5f5f5; margin: 0; }}
    .container {{ max-width: 640px; margin: 0 auto; background: white; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); padding: 32px; }}
    h3 {{ margin-top: 0; color: #333; border-bottom: 2px solid #1890ff; padding-bottom: 12px; }}
    h4 {{ margin: 0 0 16px 0; color: #1890ff; font-size: 16px; }}
    .info-section {{ background: #f0f7ff; border: 1px solid #d6e4ff; border-radius: 6px; padding: 16px; margin-bottom: 24px; }}
    .feedback-section {{ background: #fafafa; border: 1px solid #f0f0f0; border-radius: 6px; padding: 16px; margin-bottom: 24px; }}
    .feedback-item {{ border-bottom: 1px solid #f0f0f0; padding: 10px 0; }}
    .feedback-item:last-child {{ border-bottom: none; }}
    .feedback-meta {{ color: #8c8c8c; font-size: 12px; margin-bottom: 6px; }}
    .feedback-text {{ color: #262626; font-size: 14px; }}
    .feedback-empty {{ color: #8c8c8c; font-size: 13px; }}
    .info-row {{ margin-bottom: 12px; display: flex; align-items: flex-start; }}
    .info-row:last-child {{ margin-bottom: 0; }}
    .label {{ font-weight: 600; color: #595959; min-width: 80px; flex-shrink: 0; }}
    .value {{ color: #262626; flex: 1; word-break: break-word; }}
    textarea {{ width: 100%; height: 100px; padding: 12px; border: 1px solid #d9d9d9; border-radius: 4px; font-size: 14px; font-family: inherit; resize: vertical; box-sizing: border-box; }}
    textarea:focus {{ outline: none; border-color: #1890ff; box-shadow: 0 0 0 2px rgba(24,144,255,0.2); }}
    .btn-group {{ margin-top: 20px; display: flex; gap: 12px; flex-wrap: wrap; }}
    .btn {{ padding: 10px 24px; font-size: 14px; font-weight: 500; border: none; border-radius: 4px; cursor: pointer; transition: all 0.3s; }}
    .btn-false {{ background: #52c41a; color: white; }}
    .btn-false:hover {{ background: #73d13d; }}
    .btn-true {{ background: #ff4d4f; color: white; }}
    .btn-true:hover {{ background: #ff7875; }}
    .btn-restore {{ background: #1890ff; color: white; }}
    .btn-restore:hover {{ background: #40a9ff; }}
    label {{ font-weight: 600; color: #333; display: block; margin-bottom: 8px; }}
  </style>
</head>
<body>
  <div class="container">
    <h3>舆情反馈</h3>
    {info_section}
    <label>补充说明（可选）：</label>
    <textarea name="feedback_text" placeholder="例如：误报，内容与医院无关"></textarea>
    <form method="post" action="/feedback" style="margin-top: 20px;">
      <input type="hidden" name="sentiment_id" value="{sentiment_id}">
      <input type="hidden" name="sig" value="{sig}">
      <input type="hidden" name="feedback_text" id="feedback_text_hidden">
      <div class="btn-group">
        <button class="btn btn-false" type="submit" name="judgment" value="false">✓ 误判舆情</button>
        <button class="btn btn-true" type="submit" name="judgment" value="true">✗ 确认负面</button>
        {restore_button}
      </div>
    </form>
  </div>
  <script>
    document.querySelector('textarea').addEventListener('input', function() {{
      document.getElementById('feedback_text_hidden').value = this.value;
    }});
  </script>
</body>
</html>
"""


@app.post('/feedback')
def submit_feedback():
    sentiment_id = request.form.get('sentiment_id', '')
    sig = request.form.get('sig', '')
    judgment = request.form.get('judgment', '')
    action = request.form.get('action', '')
    feedback_text = (request.form.get('feedback_text') or '').strip()

    if not verify_signature(sentiment_id, sig):
        return "Invalid or expired link.", 403

    if action == 'restore':
        restore_negative_sentiment(sentiment_id)
        save_feedback({
            'sentiment_id': sentiment_id,
            'feedback_judgment': True,
            'feedback_type': 'restore_negative',
            'feedback_text': feedback_text or '恢复为负面',
            'user_id': request.remote_addr or 'web'
        })
        return "已恢复为负面，谢谢！"

    if judgment not in ('true', 'false'):
        return "Invalid feedback.", 400

    feedback_judgment = (judgment == 'true')
    feedback_type = 'true_positive' if feedback_judgment else 'false_positive'
    user_id = request.remote_addr or 'web'

    feedback_id = save_feedback({
        'sentiment_id': sentiment_id,
        'feedback_judgment': feedback_judgment,
        'feedback_type': feedback_type,
        'feedback_text': feedback_text or f"web_feedback:{feedback_type}",
        'user_id': user_id
    })

    if not feedback_judgment:
        delete_negative_sentiment(sentiment_id)
        rules = extract_rule_candidates(feedback_text)
        save_feedback_rules(feedback_id, rules, 'exclude')

    return "已收到反馈，感谢！"


if __name__ == "__main__":
    init_database()
    api_cfg = config.get("api", {})
    host = api_cfg.get("host", "0.0.0.0")
    port = int(api_cfg.get("port", 5003))
    app.run(host=host, port=port, debug=False)
