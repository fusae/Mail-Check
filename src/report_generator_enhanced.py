#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
舆情报告生成器 - 增强版
参考示例报告风格，生成更加详细和专业的舆情分析报告
Enhanced Sentiment Report Generator
"""

import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional
from collections import Counter
import re
import json
import os
import unicodedata

import requests
import yaml

try:
    import jieba
    JIEBA_AVAILABLE = True
except ImportError:
    JIEBA_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import matplotlib
    matplotlib.use('Agg')  # 使用非交互式后端
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    from matplotlib import rcParams
    plt.rcParams['axes.unicode_minus'] = False
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False


class EnhancedReportGenerator:
    """增强版舆情报告生成器"""

    def __init__(self):
        self.project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.ai_config = self._load_ai_config()
        self.chart_font_family = self._configure_matplotlib_font()

        # 创建图表输出目录
        self.charts_dir = os.path.join(self.project_root, 'data', 'charts')
        os.makedirs(self.charts_dir, exist_ok=True)

        self.platform_names = {
            '抖音': '抖音',
            'douyin': '抖音',
            '新浪微博': '微博',
            '微博': '微博',
            'weibo': '微博',
            '微信': '微信',
            'wechat': '微信',
            '新闻网站': '新闻网站',
            'news': '新闻网站',
            '黑猫投诉': '黑猫投诉',
            '今日头条': '今日头条',
            '百度贴吧': '百度贴吧',
        }

        # 情感关键词库
        self.emotion_keywords = {
            '愤怒': ['愤怒', '生气', '火大', '无语', '凭什么', '凭什么', '忍无可忍',
                    '太差了', '垃圾', '无良', '黑心', '骗子', '不负责任'],
            '悲伤': ['难过', '心痛', '悲伤', '痛苦', '不幸', '去世', '死亡', '离开',
                    '好好的一个人', '再也见不到', '遗憾', '惋惜'],
            '失望': ['失望', '失望透顶', '太失望了', '不值', '不值得', '白跑一趟',
                    '浪费时间', '不推荐', '再也不来了'],
            '质疑': ['质疑', '怀疑', '真的吗', '可信吗', '靠谱吗', '是不是',
                    '凭什么说', '证据呢', '有证据吗', '真假'],
            '担忧': ['担心', '担忧', '害怕', '恐惧', '不敢', '害怕去',
                    '有风险', '不安全', '可怕'],
        }

        # 风险等级映射
        self.risk_level_map = {
            'high': '极高',
            'medium': '高',
            'low': '中'
        }

    def _load_ai_config(self) -> Dict[str, Any]:
        """读取AI配置（DeepSeek/OpenAI兼容接口）"""
        cfg_path = os.path.join(self.project_root, "config", "config.yaml")
        try:
            with open(cfg_path, "r", encoding="utf-8") as f:
                cfg = yaml.safe_load(f) or {}
            return cfg.get("ai", {}) or {}
        except Exception:
            return {}

    def _configure_matplotlib_font(self) -> Optional[str]:
        """为服务器端图表选择可用的中文字体。"""
        if not MATPLOTLIB_AVAILABLE:
            return None

        family_candidates = [
            "Noto Sans CJK SC",
            "Noto Sans CJK JP",
            "Microsoft YaHei",
            "SimHei",
            "WenQuanYi Micro Hei",
            "PingFang SC",
            "Source Han Sans SC",
            "Arial Unicode MS",
        ]
        file_candidates = [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/arphic/ukai.ttc",
            "/usr/share/fonts/truetype/arphic/uming.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
            "C:\\Windows\\Fonts\\msyh.ttc",
            "C:\\Windows\\Fonts\\simhei.ttf",
        ]

        selected_family = None

        try:
            for font_path in file_candidates:
                if not os.path.exists(font_path):
                    continue
                fm.fontManager.addfont(font_path)
                selected_family = fm.FontProperties(fname=font_path).get_name()
                if selected_family:
                    break

            if not selected_family:
                installed_families = {font.name for font in fm.fontManager.ttflist}
                for family in family_candidates:
                    if family in installed_families:
                        selected_family = family
                        break
        except Exception:
            selected_family = None

        fallback_families = ["DejaVu Sans", "Arial", "sans-serif"]
        if selected_family:
            plt.rcParams["font.sans-serif"] = [selected_family] + fallback_families
        else:
            plt.rcParams["font.sans-serif"] = fallback_families

        plt.rcParams["axes.unicode_minus"] = False
        return selected_family

    def normalize_platform(self, platform: str) -> str:
        """标准化平台名称"""
        platform = platform.lower().strip()
        for key, value in self.platform_names.items():
            if key in platform:
                return value
        return platform

    def generate_report_data(
        self,
        df: pd.DataFrame,
        hospital_name: str,
        report_type: str = "special",
        report_period: str = None
    ) -> Dict[str, Any]:
        """
        生成增强版报告数据

        参数：
        - df: 舆情数据DataFrame
        - hospital_name: 医院名称
        - report_type: 报告类型（special/quarterly/monthly）
        - report_period: 报告周期（如"2026Q1"）
        """
        # 数据预处理
        df = self._preprocess_data(df)

        # 生成各个部分的数据
        report_data = {
            'hospital_name': hospital_name,
            'report_type': report_type,
            'report_period': report_period or self._auto_detect_period(df),
            'generated_time': datetime.now().strftime('%Y年%m月%d日 %H:%M'),
            'report_date': datetime.now().strftime('%Y年%m月%d日'),
            'report_date_range': self._get_report_date_range(df),
            'summary': self._generate_summary(df),
            'overview': self._generate_overview(df),
            'distribution': self._generate_distribution(df),
            'key_events': self._generate_key_events_enhanced(df),  # 增强版关键事件
            'sentiment': self._generate_sentiment_enhanced(df),  # 增强版情感分析
            'sentiment_analysis_new': self._generate_sentiment_analysis_new(df),  # 新增：情感分析与舆情态势
            'category_statistics': self._generate_category_statistics(df),  # 新增：舆情分类统计
            'risk_assessment': self._generate_risk_assessment_enhanced(df),  # 增强版风险评估
            'recommendations': self._generate_recommendations_enhanced(df),  # 增强版建议
            'impact_forecast': self._generate_impact_forecast(df),  # 新增：影响预测
            'spread_forecast': self._generate_spread_forecast(df),  # 新增：风险传播预测
            'response_templates': self._generate_response_templates(df),  # 新增：应对模板
            'appendix': self._generate_appendix_enhanced(df),  # 增强版附录
            'raw_dataframe': df
        }

        # 生成图表
        chart_paths = self.generate_charts(report_data, hospital_name)
        report_data['chart_paths'] = chart_paths

        return report_data

    def _preprocess_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """数据预处理"""
        df = df.copy()

        # 标准化平台名称
        df['来源_标准'] = df['来源'].apply(self.normalize_platform)

        # 解析时间
        df['创建时间_解析'] = pd.to_datetime(df['创建时间'], errors='coerce')

        # 提取日期和小时
        df['日期'] = df['创建时间_解析'].dt.date
        df['小时'] = df['创建时间_解析'].dt.hour
        df['日期_字符串'] = df['创建时间_解析'].dt.strftime('%Y-%m-%d')
        df['时间_字符串'] = df['创建时间_解析'].dt.strftime('%Y-%m-%d %H:%M')

        # 计算风险分
        df['风险分_数值'] = pd.to_numeric(df['风险分'], errors='coerce').fillna(0)

        return df

    def _auto_detect_period(self, df: pd.DataFrame) -> str:
        """自动检测报告周期"""
        if len(df) == 0:
            return datetime.now().strftime('%Y年%m月')

        dates = pd.to_datetime(df['创建时间'], errors='coerce').dropna()
        if len(dates) == 0:
            return datetime.now().strftime('%Y年%m月')

        min_date = dates.min()
        max_date = dates.max()

        if min_date.month == max_date.month:
            return min_date.strftime('%Y年%m月')
        elif min_date.year == max_date.year:
            quarter = (max_date.month - 1) // 3 + 1
            return f"{max_date.year}Q{quarter}"
        else:
            return f"{min_date.strftime('%Y年%m月')}-{max_date.strftime('%Y年%m月')}"

    def _get_report_date_range(self, df: pd.DataFrame) -> str:
        """获取报告日期范围"""
        if len(df) == 0:
            return "无数据"

        dates = pd.to_datetime(df['创建时间'], errors='coerce').dropna()
        if len(dates) == 0:
            return "无数据"

        min_date = dates.min()
        max_date = dates.max()

        return f"{min_date.strftime('%Y年%m月%d日')}-{max_date.strftime('%Y年%m月%d日')}"

    def _generate_summary(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成报告摘要（增强版）"""
        total = len(df)
        high_risk = len(df[df['严重程度'] == 'high'])
        medium_risk = len(df[df['严重程度'] == 'medium'])
        active = len(df[df['状态'] == 'active'])
        avg_risk = df['风险分_数值'].mean()

        # 估算影响人数
        estimated_reach = self._estimate_reach(df)

        # 传播峰值时间
        peak_time = self._find_peak_time(df)

        # 趋势分析
        trend = self._analyze_trend(df)

        # 危险级别判断
        danger_level = self._assess_danger_level(df)

        return {
            'total_events': total,
            'high_risk_events': high_risk,
            'medium_risk_events': medium_risk,
            'active_events': active,
            'inactive_events': total - active,
            'average_risk_score': round(avg_risk, 1),
            'estimated_reach': estimated_reach,
            'peak_time': peak_time,
            'trend': trend,
            'danger_level': danger_level,
            'platforms': df['来源_标准'].nunique(),
            'departments': df.get('科室', pd.Series()).nunique()
        }

    def _assess_danger_level(self, df: pd.DataFrame) -> str:
        """评估危险级别"""
        if len(df) == 0:
            return "无风险"

        high_risk = len(df[df['严重程度'] == 'high'])
        avg_risk = df['风险分_数值'].mean()

        if avg_risk >= 90 or high_risk >= 5:
            return "极高危险级别"
        elif avg_risk >= 70 or high_risk >= 3:
            return "高危险级别"
        elif avg_risk >= 50 or high_risk >= 1:
            return "中危险级别"
        else:
            return "低危险级别"

    def _estimate_reach(self, df: pd.DataFrame) -> str:
        """估算影响人数（增强版）"""
        if len(df) == 0:
            return "0"

        total = 0
        for _, row in df.iterrows():
            platform = row.get('来源_标准', '')
            severity = row.get('严重程度', 'low')

            # 根据平台和严重程度估算
            base_reach = 1000  # 默认1000人

            if '抖音' in platform:
                base_reach = 100000  # 抖音10万
            elif '微博' in platform:
                base_reach = 50000  # 微博5万
            elif '微信' in platform:
                base_reach = 10000  # 微信1万

            if severity == 'high':
                base_reach *= 10
            elif severity == 'medium':
                base_reach *= 3

            total += base_reach

        if total >= 100000000:
            return f"{round(total / 100000000, 1)}亿+"
        elif total >= 10000000:
            return f"{round(total / 10000000, 1)}千万+"
        elif total >= 10000:
            return f"{round(total / 10000, 1)}万+"
        elif total >= 1000:
            return f"{round(total / 1000, 1)}千+"
        else:
            return str(total)

    def _find_peak_time(self, df: pd.DataFrame) -> str:
        """找到传播峰值时间"""
        if len(df) == 0:
            return "未知"

        daily_counts = df.groupby('日期').size()
        if len(daily_counts) == 0:
            return "未知"

        peak_date = daily_counts.idxmax()
        return peak_date.strftime('%Y年%m月%d日')

    def _analyze_trend(self, df: pd.DataFrame) -> str:
        """分析趋势（增强版）"""
        if len(df) < 2:
            return "数据不足"

        df_sorted = df.sort_values('创建时间_解析')

        first_half = df_sorted[:len(df_sorted)//2]
        second_half = df_sorted[len(df_sorted)//2:]

        first_count = len(first_half)
        second_count = len(second_half)

        if second_count > first_count * 1.5:
            return "快速上升"
        elif second_count < first_count * 0.7:
            return "下降"
        else:
            return "平稳"

    def _generate_overview(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成概述数据"""
        total = len(df)

        # 按严重程度统计
        severity_counts = df['严重程度'].value_counts()

        # 按状态统计
        status_counts = df['状态'].value_counts()

        # 按平台统计
        platform_counts = df['来源_标准'].value_counts()

        return {
            'total': total,
            'severity_distribution': {
                'high': int(severity_counts.get('high', 0)),
                'medium': int(severity_counts.get('medium', 0)),
                'low': int(severity_counts.get('low', 0))
            },
            'status_distribution': status_counts.to_dict(),
            'platform_distribution': platform_counts.to_dict(),
            'average_risk_score': round(df['风险分_数值'].mean(), 1),
            'max_risk_score': int(df['风险分_数值'].max()),
            'min_risk_score': int(df['风险分_数值'].min())
        }

    def _generate_distribution(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成分布分析"""
        return {
            'time_distribution': self._analyze_time_distribution_enhanced(df),
            'platform_distribution': self._analyze_platform_distribution_enhanced(df),
            'type_distribution': self._analyze_type_distribution_enhanced(df),
            'department_distribution': self._analyze_department_distribution_enhanced(df)
        }

    def _analyze_time_distribution_enhanced(self, df: pd.DataFrame) -> Dict[str, Any]:
        """时间分布分析（增强版）"""
        if len(df) == 0:
            return {'timeline': [], 'pattern': '无数据', 'peak_hours': []}

        # 按日期统计
        daily_counts = df.groupby('日期').size().sort_index()

        # 构建时间轴
        timeline = []
        for date, count in daily_counts.items():
            date_str = date.strftime('%m月%d日')
            events = df[df['日期'] == date].sort_values('创建时间_解析')

            # 获取该日的时间段
            time_slots = []
            for _, event in events.iterrows():
                hour = event.get('小时', 0)
                time_slots.append(f"{hour:02d}:00")

            timeline.append({
                'date': date_str,
                'count': int(count),
                'time_slots': time_slots,
                'platforms': events['来源_标准'].unique().tolist()
            })

        # 按小时统计
        hourly_counts = df.groupby('小时').size()

        # 找出峰值时段
        peak_hours = sorted(hourly_counts.items(), key=lambda x: x[1], reverse=True)[:5]

        # 检测时间模式
        time_pattern = self._detect_time_pattern(df)

        return {
            'timeline': timeline,
            'daily_counts': {str(k): int(v) for k, v in daily_counts.items()},
            'hourly_counts': {int(k): int(v) for k, v in hourly_counts.items()},
            'peak_hours': [{'hour': int(h), 'count': int(c)} for h, c in peak_hours],
            'time_pattern': time_pattern
        }

    def _detect_time_pattern(self, df: pd.DataFrame) -> str:
        """检测时间模式"""
        if len(df) == 0:
            return "无数据"

        hour_counts = df.groupby('小时').size()

        # 判断是否夜间集中（22:00-02:00）
        night_hours = [22, 23, 0, 1, 2]
        night_count = sum(hour_counts.get(h, 0) for h in night_hours)

        # 判断是否工作日集中（周一至周五）
        weekday_count = len(df[df['创建时间_解析'].dt.dayofweek < 5])
        weekend_count = len(df) - weekday_count

        patterns = []
        if night_count > len(df) * 0.3:
            patterns.append("夜间集中（22:00-02:00）")
        if weekday_count > weekend_count * 2:
            patterns.append("工作日集中")
        elif weekend_count > weekday_count * 2:
            patterns.append("周末集中")

        return "、".join(patterns) if patterns else "无明显规律"

    def _analyze_platform_distribution_enhanced(self, df: pd.DataFrame) -> Dict[str, Any]:
        """平台分布分析（增强版）"""
        if len(df) == 0:
            return {'distribution': {}, 'analysis': '无数据'}

        platform_counts = df['来源_标准'].value_counts()

        distribution = {}
        for platform, count in platform_counts.items():
            percentage = (count / len(df)) * 100
            risk_level = '极高' if percentage > 70 else '高' if percentage > 30 else '中'

            # 获取该平台的风险分
            platform_df = df[df['来源_标准'] == platform]
            avg_risk = platform_df['风险分_数值'].mean()

            distribution[platform] = {
                'count': int(count),
                'percentage': round(percentage, 1),
                'risk_level': risk_level,
                'avg_risk_score': round(avg_risk, 1),
                'characteristics': self._get_platform_characteristics(platform)
            }

        # 主导平台
        dominant = platform_counts.index[0] if len(platform_counts) > 0 else "无"
        dominant_ratio = (platform_counts.iloc[0] / len(df)) * 100 if len(platform_counts) > 0 else 0

        analysis = f"主导平台：{dominant}（占{dominant_ratio:.1f}%）"

        return {
            'distribution': distribution,
            'analysis': analysis,
            'dominant_platform': dominant
        }

    def _get_platform_characteristics(self, platform: str) -> List[str]:
        """获取平台特征"""
        characteristics = {
            '抖音': ['传播速度快', '触达人群广', '情感传播强', '监管力度较弱'],
            '微博': ['话题性强', '转发传播快', '舆论发酵迅速', '媒体关注度高'],
            '微信': ['封闭传播', '圈层化明显', '长尾效应强', '难以监测'],
            '新闻网站': ['权威性高', '影响持久', '搜索引擎收录', '公信力强'],
            '黑猫投诉': ['投诉聚集地', '消费者维权', '媒体关注', '官方回复'],
            '百度贴吧': ['社群传播', '用户讨论', '长尾效应', '搜索可见']
        }
        return characteristics.get(platform, ['一般传播'])

    def _analyze_type_distribution_enhanced(self, df: pd.DataFrame) -> Dict[str, Any]:
        """类型分布分析（增强版）"""
        if len(df) == 0:
            return {'distribution': {}, 'analysis': '无数据'}

        # 假设有一个"类型"字段，如果没有则从内容推断
        if '类型' in df.columns:
            type_counts = df['类型'].value_counts()
        else:
            # 从内容推断类型
            type_counts = self._infer_event_types(df).value_counts()

        distribution = {}
        for event_type, count in type_counts.items():
            percentage = (count / len(df)) * 100
            severity = self._get_type_severity(event_type)

            distribution[event_type] = {
                'count': int(count),
                'percentage': round(percentage, 1),
                'severity': severity
            }

        return {
            'distribution': distribution,
            'total_types': len(type_counts)
        }

    def _infer_event_types(self, df: pd.DataFrame) -> pd.Series:
        """从内容推断事件类型"""
        types = []

        for _, row in df.iterrows():
            content = str(row.get('内容', '')) + str(row.get('标题', ''))

            if any(keyword in content for keyword in ['死亡', '去世', '抢救无效', '手术死亡']):
                types.append('医疗质量-死亡事件')
            elif any(keyword in content for keyword in ['投诉', '态度差', '服务差']):
                types.append('服务质量投诉')
            elif any(keyword in content for keyword in ['费用', '收费', '贵']):
                types.append('收费问题')
            elif any(keyword in content for keyword in ['等待', '排队', '时间长']):
                types.append('流程问题')
            else:
                types.append('其他')

        return pd.Series(types)

    def _get_type_severity(self, event_type: str) -> str:
        """获取类型的严重程度"""
        if '死亡' in event_type:
            return '极高'
        elif '投诉' in event_type or '费用' in event_type:
            return '高'
        else:
            return '中'

    def _analyze_department_distribution_enhanced(self, df: pd.DataFrame) -> Dict[str, Any]:
        """科室分布分析（增强版）"""
        if len(df) == 0 or '科室' not in df.columns:
            return {'distribution': {}, 'high_risk_departments': []}

        department_counts = df['科室'].value_counts()

        distribution = {}
        high_risk_departments = []

        for dept, count in department_counts.items():
            # 获取该科室的平均风险分
            dept_df = df[df['科室'] == dept]
            avg_risk = dept_df['风险分_数值'].mean()
            max_risk = dept_df['风险分_数值'].max()

            risk_level = '极高' if avg_risk >= 80 else '高' if avg_risk >= 60 else '中'

            distribution[dept] = {
                'count': int(count),
                'avg_risk_score': round(avg_risk, 1),
                'max_risk_score': int(max_risk),
                'risk_level': risk_level
            }

            if avg_risk >= 80:
                high_risk_departments.append(dept)

        return {
            'distribution': distribution,
            'high_risk_departments': high_risk_departments,
            'total_departments': len(department_counts)
        }

    def _generate_key_events_enhanced(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """生成关键事件（增强版）"""
        if len(df) == 0:
            return []

        # 按风险分排序，取前5个高风险事件
        high_risk_df = df[df['严重程度'].isin(['high', 'medium'])].sort_values('风险分_数值', ascending=False)

        # 按相似度分组（简单的标题相似度）
        event_groups = self._group_similar_events(high_risk_df)

        key_events = []
        for group_id, group_df in event_groups.items():
            if len(group_df) == 0:
                continue

            # 获取该组的代表性事件
            representative = group_df.iloc[0]

            # 构建事件脉络
            timeline = self._build_event_timeline(group_df)

            # 传播分析
            spread_analysis = self._analyze_event_spread(group_df)

            # 情感分析
            sentiment_analysis = self._analyze_event_sentiment(group_df)

            # 影响评估
            impact_assessment = self._assess_event_impact(group_df)

            # 处置建议
            recommendations = self._generate_event_recommendations(group_df)

            event = {
                'id': group_id,
                'title': representative.get('标题', '未知事件'),
                'overview': {
                    'event_time': self._extract_event_time(group_df),
                    'department': representative.get('科室', '未知'),
                    'platform': representative.get('来源_标准', '未知'),
                    'severity': representative.get('严重程度', 'unknown'),
                    'risk_score': int(representative.get('风险分_数值', 0)),
                    'total_mentions': len(group_df)
                },
                'timeline': timeline,
                'spread_analysis': spread_analysis,
                'sentiment_analysis': sentiment_analysis,
                'impact_assessment': impact_assessment,
                'recommendations': recommendations
            }

            key_events.append(event)

            if len(key_events) >= 5:
                break

        return key_events

    def _group_similar_events(self, df: pd.DataFrame) -> Dict[int, pd.DataFrame]:
        """按相似度分组事件"""
        groups = {}
        group_id = 0

        for idx, row in df.iterrows():
            title = str(row.get('标题', ''))

            # 检查是否与已有组相似
            matched = False
            for existing_id, existing_df in groups.items():
                existing_title = str(existing_df.iloc[0].get('标题', ''))
                if self._are_titles_similar(title, existing_title):
                    groups[existing_id] = pd.concat([existing_df, pd.DataFrame([row])], ignore_index=True)
                    matched = True
                    break

            if not matched:
                groups[group_id] = pd.DataFrame([row])
                group_id += 1

        return groups

    def _are_titles_similar(self, title1: str, title2: str) -> bool:
        """判断标题是否相似"""
        t1 = (title1 or "").strip()
        t2 = (title2 or "").strip()
        if not t1 or not t2:
            return False
        if t1 == t2:
            return True
        # 子串命中可视为同一事件（常见于标题前后缀变化）
        if len(t1) >= 8 and len(t2) >= 8 and (t1 in t2 or t2 in t1):
            return True

        tokens1 = self._tokenize_title(t1)
        tokens2 = self._tokenize_title(t2)
        if not tokens1 or not tokens2:
            return False

        intersection = tokens1.intersection(tokens2)
        union_size = len(tokens1 | tokens2)
        jaccard = (len(intersection) / union_size) if union_size else 0

        # 中文标题常较短，阈值适当放宽
        return len(intersection) >= 2 or jaccard >= 0.4

    def _tokenize_title(self, title: str) -> set:
        """对中文标题进行鲁棒分词，兼容无空格文本。"""
        text = (title or "").lower().strip()
        if not text:
            return set()

        tokens = set()

        # 英文/数字片段
        tokens.update(re.findall(r"[a-z0-9]+", text))

        # 中文词：优先jieba，退化到2-3字n-gram
        zh_text = "".join(re.findall(r"[\u4e00-\u9fff]", text))
        if zh_text:
            if JIEBA_AVAILABLE:
                for w in jieba.cut(zh_text):
                    w = w.strip()
                    if len(w) >= 2:
                        tokens.add(w)
            else:
                for n in (2, 3):
                    for i in range(0, max(0, len(zh_text) - n + 1)):
                        tokens.add(zh_text[i:i+n])

        stop = {"医院", "患者", "事件", "回应", "通报", "情况", "视频", "网络", "网友"}
        return {t for t in tokens if len(t) >= 2 and t not in stop}

    def _build_event_timeline(self, df: pd.DataFrame) -> Dict[str, Any]:
        """构建事件时间轴"""
        df_sorted = df.sort_values('创建时间_解析')

        stages = {
            'occurrence': [],
            'fermentation': [],
            'outbreak': [],
            'continuation': []
        }

        for _, row in df_sorted.iterrows():
            stage_info = {
                'time': row.get('时间_字符串', ''),
                'platform': row.get('来源_标准', ''),
                'description': row.get('标题', '')[:50]
            }

            # 根据时间判断阶段（简化版）
            hour = row.get('小时', 0)
            if hour < 6:
                stages['occurrence'].append(stage_info)
            elif hour < 12:
                stages['fermentation'].append(stage_info)
            elif hour < 18:
                stages['outbreak'].append(stage_info)
            else:
                stages['continuation'].append(stage_info)

        return stages

    def _analyze_event_spread(self, df: pd.DataFrame) -> Dict[str, Any]:
        """分析事件传播"""
        platforms = df['来源_标准'].value_counts().to_dict()

        # 估算传播路径
        spread_path = []
        for _, row in df.sort_values('创建时间_解析').iterrows():
            spread_path.append({
                'time': row.get('时间_字符串', ''),
                'platform': row.get('来源_标准', ''),
                'description': row.get('标题', '')[:30]
            })

        # 计算影响估算
        estimated_reach = self._estimate_reach(df)

        return {
            'platforms': platforms,
            'spread_path': spread_path,
            'estimated_reach': estimated_reach,
            'total_mentions': len(df),
            'spread_speed': self._calculate_spread_speed(df)
        }

    def _calculate_spread_speed(self, df: pd.DataFrame) -> str:
        """计算传播速度"""
        if len(df) < 2:
            return "无法计算"

        df_sorted = df.sort_values('创建时间_解析')
        time_diff = (df_sorted.iloc[-1]['创建时间_解析'] - df_sorted.iloc[0]['创建时间_解析']).total_seconds() / 3600

        if time_diff <= 0:
            return "瞬间"

        mentions_per_hour = len(df) / time_diff

        if mentions_per_hour > 10:
            return "极快（病毒式）"
        elif mentions_per_hour > 5:
            return "很快"
        elif mentions_per_hour > 1:
            return "较快"
        else:
            return "一般"

    def _analyze_event_sentiment(self, df: pd.DataFrame) -> Dict[str, Any]:
        """分析事件情感"""
        # 合并内容（优先使用“警示理由/ reason”）
        if '警示理由' in df.columns:
            all_content = ' '.join(df['警示理由'].fillna(''))
        else:
            all_content = ' '.join(df['内容'].fillna('') + ' ' + df['标题'].fillna(''))

        # 情感统计
        emotion_counts = Counter()

        if JIEBA_AVAILABLE:
            words = jieba.cut(all_content)
            word_list = list(words)

            # 统计情感词
            for emotion, keywords in self.emotion_keywords.items():
                count = sum(1 for word in word_list if word in keywords)
                if count > 0:
                    emotion_counts[emotion] += count

        # 提取高频关键词
        keywords = self._extract_keywords(all_content, top_n=20)

        # 提取公众诉求
        demands = self._extract_demands(all_content)

        return {
            'emotion_distribution': dict(emotion_counts),
            'top_keywords': keywords,
            'public_demands': demands
        }

    def _extract_keywords(self, text: str, top_n: int = 20) -> List[Dict[str, Any]]:
        """提取关键词"""
        if JIEBA_AVAILABLE:
            words = jieba.cut(text)
            word_freq = Counter(words)

            # 过滤停用词
            stop_words = {'的', '了', '是', '在', '我', '有', '和', '就', '不', '人',
                         '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去',
                         '你', '会', '着', '没有', '看', '好', '自己', '这', '但'}

            filtered = {k: v for k, v in word_freq.items()
                       if len(k) > 1 and k not in stop_words and v > 1}

            top_keywords = sorted(filtered.items(), key=lambda x: x[1], reverse=True)[:top_n]

            return [{'keyword': k, 'count': v} for k, v in top_keywords]
        else:
            return []

    def _extract_demands(self, text: str) -> List[str]:
        """提取公众诉求"""
        demand_patterns = [
            (r'要求.*?责任', '要求医院承担责任'),
            (r'(要求|请求).*?调查', '要求调查事件真相'),
            (r'(要求|请求).*?道歉', '要求道歉'),
            (r'(要求|请求).*?赔偿', '要求赔偿'),
            (r'(要求|请求).*?退款', '要求退款'),
            (r'(要求|请求).*?公开', '要求公开信息'),
            (r'(要求|请求).*?处理', '要求处理相关人员'),
        ]

        demands = []
        for pattern, demand in demand_patterns:
            if re.search(pattern, text):
                demands.append(demand)

        return demands

    def _assess_event_impact(self, df: pd.DataFrame) -> Dict[str, Any]:
        """评估事件影响"""
        avg_risk = df['风险分_数值'].mean()
        max_risk = df['风险分_数值'].max()
        total_mentions = len(df)

        # 社会影响
        social_impact = []
        if avg_risk >= 80:
            social_impact.append("严重损害医院声誉")
        if avg_risk >= 70:
            social_impact.append("引发公众对医院水平的质疑")
        if total_mentions > 5:
            social_impact.append("可能引发媒体跟进报道")

        # 潜在风险
        potential_risks = []
        if '死亡' in ' '.join(df['内容'].fillna('')):
            potential_risks.append("可能引发法律诉讼")
            potential_risks.append("可能影响医院评级")
        if avg_risk >= 70:
            potential_risks.append("可能导致其他患者流失")

        return {
            'social_impact': social_impact,
            'potential_risks': potential_risks,
            'legal_risk': '高' if '死亡' in ' '.join(df['内容'].fillna('')) else '中',
            'media_risk': '高' if total_mentions > 5 else '中'
        }

    def _generate_event_recommendations(self, df: pd.DataFrame) -> Dict[str, List[str]]:
        """生成事件处置建议"""
        avg_risk = df['风险分_数值'].mean()

        immediate = []
        short_term = []
        long_term = []

        if avg_risk >= 80:
            immediate.extend([
                "立即启动危机公关响应",
                "发布官方声明",
                "启动内部调查",
                "主动与相关方沟通"
            ])

        if avg_risk >= 60:
            immediate.append("密切监控舆情发展")
            short_term.extend([
                "准备媒体应对材料",
                "评估法律风险"
            ])

        long_term.extend([
            "改进相关医疗流程",
            "加强医患沟通培训",
            "建立危机预警机制"
        ])

        return {
            'immediate': immediate,
            'short_term': short_term,
            'long_term': long_term
        }

    def _extract_event_time(self, df: pd.DataFrame) -> str:
        """提取事件时间"""
        if len(df) == 0:
            return "未知"

        min_time = df['创建时间_解析'].min()
        max_time = df['创建时间_解析'].max()

        if min_time == max_time:
            return min_time.strftime('%Y年%m月%d日')
        else:
            return f"{min_time.strftime('%m月%d日')}-{max_time.strftime('%m月%d日')}"

    def _generate_sentiment_enhanced(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成情感分析（增强版）"""
        if len(df) == 0:
            return {'emotion_distribution': {}, 'top_keywords': [], 'public_demands': []}

        # 合并所有内容（优先使用标题+正文，避免关键词被“AI理由模板词”污染）
        if '内容' in df.columns or '标题' in df.columns:
            all_content = ' '.join(df.get('标题', pd.Series()).fillna('') + ' ' + df.get('内容', pd.Series()).fillna(''))
            if not all_content.strip() and '警示理由' in df.columns:
                all_content = ' '.join(df['警示理由'].fillna(''))
        elif '警示理由' in df.columns:
            all_content = ' '.join(df['警示理由'].fillna(''))
        else:
            all_content = ""

        # 情感统计
        emotion_counts = Counter()

        if JIEBA_AVAILABLE:
            words = jieba.cut(all_content)
            word_list = list(words)

            for emotion, keywords in self.emotion_keywords.items():
                count = sum(1 for word in word_list if word in keywords)
                if count > 0:
                    emotion_counts[emotion] = count

        # 提取关键词
        keywords = self._extract_keywords(all_content, top_n=30)

        # 提取诉求
        demands = self._extract_demands(all_content)

        return {
            'emotion_distribution': dict(emotion_counts),
            'top_keywords': keywords,
            'public_demands': demands
        }

    def _generate_risk_assessment_enhanced(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成风险评估（增强版）"""
        if len(df) == 0:
            return {'current_risks': [], 'risk_levels': {}}

        current_risks = []
        risk_levels = {
            'red': [],
            'orange': [],
            'yellow': []
        }

        # 按科室分析风险
        if '科室' in df.columns:
            for dept in df['科室'].unique():
                dept_df = df[df['科室'] == dept]
                avg_risk = dept_df['风险分_数值'].mean()
                max_risk = dept_df['风险分_数值'].max()
                count = len(dept_df)

                if avg_risk >= 80:
                    level = 'red'
                    level_text = '红色预警（极高风险）'
                elif avg_risk >= 60:
                    level = 'orange'
                    level_text = '橙色预警（高风险）'
                else:
                    level = 'yellow'
                    level_text = '黄色预警（中风险）'

                risk_info = {
                    'department': dept,
                    'avg_risk_score': round(avg_risk, 1),
                    'max_risk_score': int(max_risk),
                    'event_count': int(count),
                    'level_text': level_text
                }

                current_risks.append(risk_info)
                risk_levels[level].append(dept)

        # 按事件类型分析
        event_type_risks = []
        for event_type in df.get('类型', pd.Series()).unique():
            if pd.isna(event_type):
                continue
            type_df = df[df['类型'] == event_type]
            avg_risk = type_df['风险分_数值'].mean()

            event_type_risks.append({
                'type': event_type,
                'avg_risk_score': round(avg_risk, 1),
                'event_count': len(type_df)
            })

        return {
            'current_risks': sorted(current_risks, key=lambda x: x['avg_risk_score'], reverse=True),
            'risk_levels': risk_levels,
            'event_type_risks': event_type_risks,
            'overall_risk_level': self._assess_danger_level(df)
        }

    def _generate_recommendations_enhanced(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成应对建议（增强版）"""
        ai_recs = self._generate_ai_recommendations(df)
        if ai_recs:
            return ai_recs

        avg_risk = df['风险分_数值'].mean() if len(df) > 0 else 0

        immediate = []
        short_term = []
        long_term = []

        # 立即措施（24小时内）
        if avg_risk >= 80:
            immediate.extend([
                "立即启动危机公关响应",
                "发布官方声明",
                "启动内部调查",
                "准备法律应对",
                "主动与相关方沟通"
            ])

        # 短期措施（1周内）
        if avg_risk >= 60:
            short_term.extend([
                "公布调查结果",
                "处理相关责任人",
                "整改医疗流程",
                "加强医患沟通培训",
                "建立危机预警机制"
            ])

        # 长期措施（1-3个月）
        long_term.extend([
            "🏥 提高医疗质量",
            "💬 改善服务态度",
            "⚡ 优化服务流程",
            "🔒 加强危机预防",
            "📖 建立投诉处理制度"
        ])

        # 重点防控方向
        prevention = {
            'short_term': [
                "高风险科室专项整治",
                "危机管理机制建设",
                "全院服务质量提升"
            ],
            'medium_term': [
                "医患沟通培训",
                "服务流程优化",
                "投诉处理机制"
            ]
        }

        # 舆情监测重点
        monitoring = {
            'keywords': self._generate_monitoring_keywords(df),
            'platforms': list(df['来源_标准'].unique()),
            'frequency': '实时监测（7x24小时）'
        }

        return {
            'immediate_actions': immediate,
            'short_term_actions': short_term,
            'long_term_actions': long_term,
            'prevention': prevention,
            'monitoring': monitoring,
            'generation_source': 'rule',
            'ai_summary': ''
        }

    def _generate_ai_recommendations(self, df: pd.DataFrame) -> Optional[Dict[str, Any]]:
        """调用AI生成处置建议，失败时返回None并由规则回退。"""
        if df is None or len(df) == 0:
            return None

        api_key = (self.ai_config.get("api_key") or "").strip()
        api_url = (self.ai_config.get("api_url") or "").strip()
        model = (self.ai_config.get("model") or "").strip()
        if not api_key or not api_url or not model:
            return None

        sample_df = df.copy()
        if '风险分_数值' in sample_df.columns:
            sample_df = sample_df.sort_values(by=['风险分_数值', '创建时间_解析'], ascending=[False, False], na_position='last')
        sample_df = sample_df.head(25)

        lines = []
        for idx, row in enumerate(sample_df.itertuples(index=False), 1):
            title = (getattr(row, '标题', '') or '无标题').replace('\n', ' ').strip()
            source = (getattr(row, '来源_标准', '') or getattr(row, '来源', '') or '未知').strip()
            risk = int(getattr(row, '风险分_数值', 0) or 0)
            created = (getattr(row, '时间_字符串', '') or str(getattr(row, '创建时间', '') or '')[:19]).strip()
            reason = (getattr(row, '警示理由', '') or '').replace('\n', ' ').strip()
            content = (getattr(row, '内容', '') or '').replace('\n', ' ').strip()
            url = (getattr(row, '原文链接', '') or '').strip()
            body = content[:160] if content else reason[:160]
            lines.append(
                f"{idx}. 时间:{created} 平台:{source} 风险:{risk} 标题:{title} 摘要:{body} 链接:{url or '无'}"
            )

        hospital_name = str(sample_df.get('医院', pd.Series()).iloc[0] if '医院' in sample_df.columns and len(sample_df) > 0 else "该医院")
        total_events = int(len(df))
        high_risk = int(len(df[df['严重程度'] == 'high'])) if '严重程度' in df.columns else 0

        user_prompt = (
            f"你是医院舆情处置专家。请基于以下舆情为{hospital_name}输出可执行处置建议。\n"
            f"数据概览：总条数{total_events}，高风险{high_risk}。\n"
            "要求：\n"
            "1) 仅输出JSON，不要Markdown。\n"
            "2) 字段必须包含：\n"
            "{\n"
            "  \"ai_summary\": \"一句话总体判断(<=80字)\",\n"
            "  \"key_risks\": [\"风险点1\", \"风险点2\", \"风险点3\"],\n"
            "  \"immediate_actions\": [\"24小时内动作...\"],\n"
            "  \"short_term_actions\": [\"72小时内动作...\"],\n"
            "  \"long_term_actions\": [\"一月内机制动作...\"]\n"
            "}\n"
            "3) 每个数组3-6条，避免空泛话术，尽量具体到责任动作。\n\n"
            "舆情样本：\n" + "\n".join(lines)
        )

        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        }
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": "你是专业的医疗舆情应对顾问。"},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": float(self.ai_config.get("temperature", 0.3)),
            "max_tokens": int(self.ai_config.get("max_tokens", 800)),
        }

        try:
            resp = requests.post(
                api_url,
                headers=headers,
                json=payload,
                timeout=int(self.ai_config.get("timeout", 45))
            )
            resp.raise_for_status()
            result = resp.json()
            content = result["choices"][0]["message"]["content"]
            parsed = self._parse_ai_json(content)
            if not parsed:
                return None

            immediate = self._to_clean_list(parsed.get("immediate_actions"))
            short_term = self._to_clean_list(parsed.get("short_term_actions"))
            long_term = self._to_clean_list(parsed.get("long_term_actions"))
            key_risks = self._to_clean_list(parsed.get("key_risks"))
            summary = str(parsed.get("ai_summary") or "").strip()
            if not immediate and not short_term and not long_term:
                return None

            return {
                "immediate_actions": immediate,
                "short_term_actions": short_term,
                "long_term_actions": long_term,
                "prevention": {"short_term": [], "medium_term": []},
                "monitoring": {
                    "keywords": self._generate_monitoring_keywords(df),
                    "platforms": list(df['来源_标准'].unique()) if '来源_标准' in df.columns else [],
                    "frequency": "实时监测（7x24小时）",
                },
                "generation_source": "ai",
                "ai_summary": summary,
                "key_risks": key_risks
            }
        except Exception:
            return None

    def _parse_ai_json(self, text: str) -> Optional[Dict[str, Any]]:
        """容错解析AI返回JSON。"""
        if not text:
            return None
        raw = text.strip()
        try:
            return json.loads(raw)
        except Exception:
            pass
        start = raw.find("{")
        end = raw.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return None
        try:
            return json.loads(raw[start:end + 1])
        except Exception:
            return None

    def _to_clean_list(self, value: Any) -> List[str]:
        if not isinstance(value, list):
            return []
        out = []
        for item in value:
            s = str(item or "").strip()
            if s:
                out.append(s)
        return out

    def _sanitize_markdown_text(self, text: str) -> str:
        """移除Markdown中的emoji和视觉类特殊符号，保留Markdown结构。"""
        if not text:
            return ""

        cleaned = text.replace("\ufe0f", "").replace("\u200d", "")
        out: List[str] = []

        for ch in cleaned:
            codepoint = ord(ch)
            category = unicodedata.category(ch)

            if (
                0x1F000 <= codepoint <= 0x1FAFF
                or 0x2600 <= codepoint <= 0x27BF
                or 0x2B00 <= codepoint <= 0x2BFF
                or 0x2190 <= codepoint <= 0x21FF
                or category in {"So", "Sk", "Cs"}
            ):
                continue

            out.append(ch)

        return "".join(out)

    def _generate_monitoring_keywords(self, df: pd.DataFrame) -> List[str]:
        """生成监测关键词"""
        if '医院' in df.columns:
            hospital_names = df['医院'].unique().tolist()
        else:
            hospital_names = []

        keywords = []
        for name in hospital_names[:5]:
            keywords.extend([
                name,
                f"{name} 死亡",
                f"{name} 投诉",
                f"{name} 手术"
            ])

        return keywords[:20]

    def _generate_impact_forecast(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成影响预测"""
        avg_risk = df['风险分_数值'].mean() if len(df) > 0 else 0

        # 短期影响（1-7天）
        short_term = []
        if avg_risk >= 70:
            short_term.extend([
                "平台持续发酵",
                "可能出现更多相关内容",
                "医院网络评分下降"
            ])

        # 中期影响（1-4周）
        medium_term = []
        if avg_risk >= 60:
            medium_term.extend([
                "传统媒体可能报道",
                "监管部门可能关注",
                "可能引发法律诉讼"
            ])

        if avg_risk >= 70:
            medium_term.append("就诊量可能下降10-20%")

        # 长期影响（1-3个月）
        long_term = []
        if avg_risk >= 50:
            long_term.extend([
                "医院声誉受损",
                "品牌形象下降",
                "市场份额流失"
            ])

        # 法律风险评估
        legal_risk = {
            'probability': '80%' if '死亡' in ' '.join(df['内容'].fillna('')) else '30%',
            'estimated_amount': '50-200万' if avg_risk >= 70 else '10-50万',
            'description': '医疗损害赔偿诉讼风险较高' if avg_risk >= 70 else '存在诉讼风险'
        }

        return {
            'short_term': short_term,
            'medium_term': medium_term,
            'long_term': long_term,
            'legal_risk': legal_risk
        }

    def _generate_response_templates(self, df: pd.DataFrame) -> Dict[str, str]:
        """生成应对模板"""
        hospital_name = df.get('医院', pd.Series()).iloc[0] if len(df) > 0 and '医院' in df.columns else "我院"

        # 首次回应模板
        first_response = f"""关于网传{hospital_name}患者事件的首次回应

我院关注到网络平台出现关于我院的舆情，对此我们深表关切。
医院已第一时间成立专项调查组，对事件进行全面调查。

我们承诺：
1. 秉持客观、公正、透明的原则
2. 尽快查明事实真相
3. 依法依规处理
4. 及时向社会公布调查进展

感谢社会各界监督。

{hospital_name}
{datetime.now().strftime('%Y年%m月%d日')}
"""

        # 调查进展模板
        progress_update = f"""关于患者事件调查进展的通报

自启动调查以来，我院已完成以下工作：

一、已完成：
1. 封存全部病历资料
2. 调阅相关监控录像
3. 约谈相关医护人员
4. 与相关方取得联系

二、正在进行：
1. 医疗过程评估
2. 病历资料分析
3. 专家论证
4. 责任认定

三、后续安排：
1. 尽快公布调查结果
2. 依法依规处理
3. 改进医疗服务

感谢社会各界的关心和监督。

{hospital_name}
{datetime.now().strftime('%Y年%m月%d日')}
"""

        return {
            'first_response': first_response,
            'progress_update': progress_update
        }

    def _generate_appendix_enhanced(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成附录（增强版）"""
        if len(df) == 0:
            return {'event_list': [], 'contact_info': {}}

        # 完整事件清单
        event_list = []
        for _, row in df.sort_values('创建时间_解析', ascending=False).iterrows():
            event_list.append({
                'id': row.get('ID', ''),
                'time': row.get('时间_字符串', ''),
                'platform': row.get('来源_标准', ''),
                'type': row.get('类型', '未知'),
                'department': row.get('科室', ''),
                'risk_score': int(row.get('风险分_数值', 0)),
                'status': row.get('状态', 'unknown'),
                'title': row.get('标题', '')[:50]
            })

        # 传播路径（简化版）
        spread_path = self._build_spread_path(df)

        # 联系方式
        contact_info = {
            'monitoring_center': {
                '负责人': '[填写]',
                '电话': '[填写]',
                '邮箱': '[填写]'
            },
            'crisis_team': {
                '组长': '院长',
                '成员': ['医务科', '宣传科', '法务科']
            }
        }

        return {
            'event_list': event_list,
            'spread_path': spread_path,
            'contact_info': contact_info
        }

    def _build_spread_path(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """构建传播路径"""
        if len(df) == 0:
            return []

        df_sorted = df.sort_values('创建时间_解析')

        path = []
        for _, row in df_sorted.iterrows():
            path.append({
                'time': row.get('时间_字符串', ''),
                'platform': row.get('来源_标准', ''),
                'title': row.get('标题', '')[:30],
                'description': row.get('内容', '')[:50]
            })

        return path[:50]  # 最多显示50条

    def generate_markdown_report(self, report_data: Dict[str, Any]) -> str:
        """生成Markdown格式报告（简洁版）"""
        lines = []

        # 封面
        lines.append(f"# {report_data['hospital_name']}舆情处置简报\n")
        lines.append(f"**统计周期：** {report_data['report_date_range']}")
        lines.append(f"**生成时间：** {report_data['generated_time']}")
        lines.append(f"**报告类型：** 舆情监测分析报告\n")
        lines.append("---\n")

        lines.extend(self._format_exec_summary_section(report_data))  # 1. 核心结论
        lines.extend(self._format_spread_snapshot_section(report_data))  # 3→2. 传播快照
        lines.extend(self._format_keyword_cloud_section(report_data))  # 4→3. 关键词云
        lines.extend(self._format_sentiment_analysis_section(report_data))  # 6→4. 情感分析
        lines.extend(self._format_category_statistics_section(report_data))  # 7→5. 分类统计
        lines.extend(self._format_top_events_section(report_data))  # 2→6. 重点事件
        lines.extend(self._format_action_plan_section(report_data))  # 5→7. 处置动作
        lines.extend(self._format_spread_forecast_section(report_data))  # 8. 传播预测
        lines.extend(self._format_compact_appendix_section(report_data))  # 9. 附录

        lines.append("\n---")
        lines.append("\n> 注：本简报用于快速决策，建议配合原文链接进行人工复核。\n")
        return self._sanitize_markdown_text('\n'.join(lines))

    def _format_exec_summary_section(self, data: Dict[str, Any]) -> List[str]:
        lines = []
        summary = data.get('summary', {})
        overview = data.get('overview', {})
        platforms = list((overview.get('platform_distribution') or {}).keys())

        lines.append("## 一、核心结论\n")

        danger_level = summary.get('danger_level', '未知')
        lines.append(f"### 风险等级\n")
        lines.append(f"**{danger_level}**\n")

        lines.append(f"### 关键数据\n")
        lines.append(f"- 负面舆情总数：{summary.get('total_events', 0)} 条")
        lines.append(f"  - 高危事件：{summary.get('high_risk_events', 0)} 条")
        lines.append(f"  - 中危事件：{summary.get('medium_risk_events', 0)} 条")
        lines.append(f"- 传播峰值：{summary.get('peak_time', '未知')}")
        lines.append(f"- 发展趋势：{summary.get('trend', '未知')}")
        if platforms:
            lines.append(f"- 主要平台：{', '.join(platforms[:3])}")
        lines.append("")
        return lines

    def _format_top_events_section(self, data: Dict[str, Any]) -> List[str]:
        lines = []
        df = data.get('raw_dataframe', pd.DataFrame())
        if df is None or len(df) == 0:
            return lines
        lines.append("## 六、重点事件 Top 10（按风险与时间排序）\n")
        work = df.copy()
        work['风险分_数值'] = pd.to_numeric(work.get('风险分_数值', 0), errors='coerce').fillna(0)
        work = work.sort_values(by=['风险分_数值', '创建时间_解析'], ascending=[False, False], na_position='last')
        topn = work.head(10)
        lines.append("| 序号 | 标题 | 来源 | 风险分 | 关键判断 | 原文 |")
        lines.append("|---|---|---|---:|---|---|")
        for idx, row in enumerate(topn.itertuples(index=False), 1):
            title = getattr(row, '标题', '') or '无标题'
            source = getattr(row, '来源_标准', '') or getattr(row, '来源', '') or '未知'
            score = int(getattr(row, '风险分_数值', 0) or 0)
            reason = (getattr(row, '警示理由', '') or '').replace('\n', ' ').strip()
            reason = reason[:40] + ('...' if len(reason) > 40 else '')
            url = getattr(row, '原文链接', '') or ''
            url_col = f"[查看]({url})" if url else "无"
            lines.append(f"| {idx} | {title[:40]}{'...' if len(title)>40 else ''} | {source} | {score} | {reason or '无'} | {url_col} |")
        lines.append("")
        return lines

    def _format_spread_snapshot_section(self, data: Dict[str, Any]) -> List[str]:
        lines = []
        dist = data.get('distribution', {})
        lines.append("## 二、传播快照\n")

        # 插入平台分布饼图
        chart_paths = data.get('chart_paths', {})
        if chart_paths.get('platform_pie'):
            lines.append(f"![平台分布]({chart_paths['platform_pie']})\n")

        lines.append("### 传播时段分析\n")
        time_dist = dist.get('time_distribution', {}) or {}
        peak_hours = time_dist.get('peak_hours', [])[:3]
        if peak_hours:
            peak_text = "、".join([f"{h.get('hour', 0):02d}:00({h.get('count', 0)}条)" for h in peak_hours])
            lines.append(f"- 峰值时段：{peak_text}")
        pattern = time_dist.get('time_pattern')
        if pattern:
            lines.append(f"- 时间规律：{pattern}")

        lines.append("\n### 平台分布详情\n")
        p_dist = (dist.get('platform_distribution', {}) or {}).get('distribution', {}) or {}
        if p_dist:
            top_p = sorted(p_dist.items(), key=lambda kv: kv[1].get('count', 0), reverse=True)[:5]
            for name, info in top_p:
                lines.append(f"- {name}：{info.get('count', 0)}条 ({info.get('percentage', 0)}%)")
        lines.append("")
        return lines

    def _format_action_plan_section(self, data: Dict[str, Any]) -> List[str]:
        lines = []
        recs = data.get('recommendations', {}) or {}
        lines.append("## 七、处置动作（可执行）\n")
        ai_summary = (recs.get('ai_summary') or '').strip()
        if ai_summary:
            lines.append(f"- AI判断：{ai_summary}")
        key_risks = recs.get('key_risks', []) or []
        if key_risks:
            lines.append(f"- 重点风险：{'；'.join([str(x) for x in key_risks[:3]])}")
        immediate = recs.get('immediate_actions', [])[:5]
        short_term = recs.get('short_term_actions', [])[:5]
        long_term = recs.get('long_term_actions', [])[:5]
        if immediate:
            lines.append("### 24小时内")
            for item in immediate:
                lines.append(f"- {item}")
        if short_term:
            lines.append("\n### 72小时内")
            for item in short_term:
                lines.append(f"- {item}")
        lines.append("\n### 本周内")
        if long_term:
            for item in long_term:
                lines.append(f"- {item}")
        else:
            lines.append("- 复核 Top 10 事件原文，标记误报并沉淀规则")
            lines.append("- 对高频平台（如抖音）建立重点监控词与人工复查机制")
        lines.append("")
        return lines

    def _format_keyword_cloud_section(self, data: Dict[str, Any]) -> List[str]:
        lines = []
        sentiment = data.get('sentiment', {}) or {}
        image = sentiment.get('wordcloud_image')
        keywords = sentiment.get('top_keywords', []) or []

        lines.append("## 三、关键词云与高频词\n")

        lines.append("### 关键词云图\n")
        if image:
            lines.append(f"![关键词云图]({image})\n")
        else:
            lines.append("关键词云图未生成\n")

        lines.append("### 高频词 Top 15\n")
        if keywords:
            lines.append("| 排名 | 关键词 | 出现次数 |")
            lines.append("|:----:|--------|:--------:|")
            for idx, item in enumerate(keywords[:15], 1):
                if isinstance(item, dict):
                    kw = item.get('keyword', '')
                    cnt = item.get('count', 0)
                else:
                    try:
                        kw, cnt = item
                    except Exception:
                        continue
                if not kw:
                    continue
                lines.append(f"| {idx} | {kw} | {cnt}次 |")
        else:
            lines.append("无可用高频词")
        lines.append("")
        return lines

    def _format_compact_appendix_section(self, data: Dict[str, Any]) -> List[str]:
        lines = []
        appendix = data.get('appendix', {}) or {}
        event_list = appendix.get('event_list', [])[:20]
        lines.append("## 九、附录（事件清单）\n")
        if not event_list:
            lines.append("- 无")
            lines.append("")
            return lines
        lines.append("| 时间 | 平台 | 风险分 | 状态 | 标题 |")
        lines.append("|---|---|---:|---|---|")
        for event in event_list:
            t = (event.get('time') or '')[:16]
            p = event.get('platform') or '未知'
            s = event.get('risk_score') or 0
            st = event.get('status') or 'unknown'
            title = event.get('title') or '无标题'
            lines.append(f"| {t} | {p} | {s} | {st} | {title[:50]} |")
        lines.append("")
        return lines

    def generate_word_report(self, report_data: Dict[str, Any], output_path: str) -> None:
        """生成Word格式报告（基于Markdown渲染为简化版排版）"""
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装python-docx: pip install python-docx")

        markdown_text = self.generate_markdown_report(report_data)
        doc = Document()
        self._render_markdown_to_docx(doc, markdown_text)
        doc.save(output_path)

    def _render_markdown_to_docx(self, doc: Document, markdown_text: str) -> None:
        """将Markdown文本渲染为基础Word格式"""
        lines = markdown_text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # 跳过分隔线
            if line.strip() in ("---", "----", "-----"):
                i += 1
                continue

            # 代码块
            if line.strip().startswith("```"):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i].rstrip())
                    i += 1
                if code_lines:
                    doc.add_paragraph("\n".join(code_lines))
                i += 1
                continue

            # 图片
            if line.strip().startswith("!["):
                # 提取图片路径 ![alt](path)
                import re
                match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line.strip())
                if match:
                    alt_text = match.group(1)
                    img_path = match.group(2)
                    try:
                        # 添加图片到Word文档
                        if os.path.exists(img_path):
                            doc.add_picture(img_path, width=Inches(5.5))
                            # 添加图片说明
                            if alt_text:
                                p = doc.add_paragraph(alt_text)
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        # 如果图片加载失败，添加文本说明
                        doc.add_paragraph(f"[图片: {alt_text}]")
                i += 1
                continue

            # 表格
            if line.strip().startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1

                rows = []
                for row_line in table_lines:
                    parts = [p.strip() for p in row_line.strip("|").split("|")]
                    rows.append(parts)

                # 跳过对齐行
                if len(rows) > 1 and all(set(cell) <= set("-:") for cell in rows[1]):
                    rows.pop(1)

                if rows:
                    cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=1, cols=cols)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for idx, value in enumerate(rows[0]):
                        self._add_formatted_text(hdr_cells[idx].paragraphs[0], value, bold=True)
                    for row in rows[1:]:
                        row_cells = table.add_row().cells
                        for idx, value in enumerate(row):
                            self._add_formatted_text(row_cells[idx].paragraphs[0], value)
                continue

            # 标题
            if line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
                i += 1
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
                i += 1
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
                i += 1
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
                i += 1
                continue

            # 列表
            if line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, line[2:])
                i += 1
                continue

            # 普通文本
            if line.strip():
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
            i += 1

    def _add_formatted_text(self, paragraph, text: str, bold: bool = False) -> None:
        """添加带格式的文本到段落（处理粗体等markdown格式）"""
        import re
        # 处理粗体 **text**
        parts = re.split(r'(\*\*[^\*]+\*\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 粗体文本
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part:
                # 普通文本
                run = paragraph.add_run(part)
                if bold:
                    run.bold = True

    def _format_summary_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化概述部分"""
        lines = []
        summary = data['summary']
        scope = data.get('data_scope', {}) or {}

        lines.append("## 一、报告概述\n")
        lines.append("### 1.1 舆情总体态势\n")
        lines.append(f"**{summary['danger_level']}**\n")

        lines.append("### 1.2 关键数据摘要\n")
        lines.append("| 指标 | 数值 | 说明 |")
        lines.append("|------|------|------|")
        lines.append(f"| **负面舆情总数** | {summary['total_events']}条 | 全部为{'高风险' if summary['high_risk_events'] > 0 else '中低风险'} |")
        lines.append(f"| **高危事件数量** | {summary['high_risk_events']}起 | 需重点关注 |")
        lines.append(f"| **影响人数估算** | {summary['estimated_reach']} | 仅估算 |")
        lines.append(f"| **传播峰值时间** | {summary['peak_time']} | 集中爆发期 |")
        lines.append(f"| **涉及平台** | {summary['platforms']}个 | {', '.join(list(data.get('overview', {}).get('platform_distribution', {}).keys())[:3])} |")
        lines.append(f"| **平均风险分** | {summary['average_risk_score']}分 | 满分100分 |\n")

        if scope:
            include_dismissed = "是" if scope.get("include_dismissed") else "否"
            dedupe_by_event = "是" if scope.get("dedupe_by_event") else "否"
            lines.append("### 1.3 数据口径\n")
            lines.append("| 口径项 | 值 |")
            lines.append("|------|----|")
            lines.append(f"| 包含误报数据 | {include_dismissed} |")
            lines.append(f"| 按事件归并去重 | {dedupe_by_event} |")
            lines.append(f"| 原始记录数 | {scope.get('raw_count', summary['total_events'])} |")
            lines.append(f"| 报告纳入数 | {scope.get('included_count', summary['total_events'])} |")
            lines.append(f"| 去重/过滤数 | {scope.get('excluded_count', 0)} |\n")

        lines.append("### 1.4 环比变化\n")
        lines.append(f"- {summary['trend']}\n")
        lines.append("---\n")

        return lines

    def _format_distribution_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化分布分析部分"""
        lines = []
        dist = data['distribution']

        lines.append("## 二、舆情分布分析\n")

        # 时间分布
        lines.append("### 2.1 时间分布\n")
        time_dist = dist.get('time_distribution', {})
        timeline = time_dist.get('timeline', [])

        if timeline:
            lines.append("**时间轴：**\n")
            lines.append("```")
            for item in timeline[:10]:
                lines.append(f"{item['date']}: {item['count']}条")
                if item.get('time_slots'):
                    lines.append(f"  时段: {', '.join(item['time_slots'][:3])}")
            lines.append("```")

        lines.append(f"\n**关键发现：**")
        lines.append(f"- {time_dist.get('time_pattern', '无明显规律')}")
        peak_hours_str = ', '.join([f"{h['hour']}:00" for h in time_dist.get('peak_hours', [])[:3]])
        lines.append(f"- 峰值时段: {peak_hours_str}\n")

        # 平台分布
        lines.append("### 2.2 平台分布\n")
        platform_dist = dist.get('platform_distribution', {})
        platforms = platform_dist.get('distribution', {})

        if platforms:
            lines.append("| 平台 | 数量 | 占比 | 风险等级 |")
            lines.append("|------|------|------|----------|")
            for platform, info in platforms.items():
                lines.append(f"| **{platform}** | {info['count']}条 | {info['percentage']}% | {info['risk_level']} |")

        lines.append(f"\n**{platform_dist.get('analysis', '')}**\n")

        # 类型分布
        lines.append("### 2.3 类型分布\n")
        type_dist = dist.get('type_distribution', {}).get('distribution', {})

        if type_dist:
            lines.append("| 类型 | 数量 | 占比 | 严重程度 |")
            lines.append("|------|------|------|----------|")
            for event_type, info in type_dist.items():
                lines.append(f"| **{event_type}** | {info['count']}条 | {info['percentage']}% | {info['severity']} |\n")

        # 科室分布
        if '科室' in data.get('raw_dataframe', pd.DataFrame()).columns:
            lines.append("### 2.4 科室分布\n")
            dept_dist = dist.get('department_distribution', {}).get('distribution', {})

            if dept_dist:
                lines.append("| 科室 | 事件数 | 风险等级 | 平均风险分 |")
                lines.append("|------|--------|----------|------------|")
                for dept, info in list(dept_dist.items())[:5]:
                    lines.append(f"| **{dept}** | {info['count']} | {info['risk_level']} | {info['avg_risk_score']} |")

                high_risk = dist.get('department_distribution', {}).get('high_risk_departments', [])
                if high_risk:
                    lines.append(f"\n**重点监控科室：** {', '.join([f'**{d}**' for d in high_risk])}\n")

        lines.append("---\n")

        return lines

    def _format_key_events_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化关键事件部分"""
        lines = []
        events = data.get('key_events', [])

        if not events:
            return []

        lines.append("## 三、重点负面事件详析\n")

        for idx, event in enumerate(events[:3], 1):
            risk_level = '【高风险】' if event['overview']['severity'] == 'high' else '【中风险】'
            lines.append(f"### {risk_level} 事件{idx}：{event['title'][:50]}\n")

            # 概况
            lines.append("#### 3.{}.1 事件概况".format(idx))
            lines.append("| 项目 | 详情 |")
            lines.append("|------|------|")
            lines.append(f"| **事件时间** | {event['overview']['event_time']} |")
            lines.append(f"| **涉事科室** | {event['overview']['department']} |")
            lines.append(f"| **首发平台** | {event['overview']['platform']} |")
            lines.append(f"| **风险评分** | {event['overview']['risk_score']}/100 |")
            lines.append(f"| **传播次数** | {event['overview']['total_mentions']}次 |\n")

            # 传播分析
            if event.get('spread_analysis'):
                lines.append("#### 3.{}.2 传播分析".format(idx))
                spread = event['spread_analysis']
                lines.append(f"- **传播速度：** {spread.get('spread_speed', '未知')}")
                lines.append(f"- **影响估算：** {spread.get('estimated_reach', '未知')}")
                lines.append(f"- **涉及平台：** {', '.join(spread.get('platforms', {}).keys())}\n")

            # 情感分析
            if event.get('sentiment_analysis'):
                lines.append("#### 3.{}.3 情感倾向".format(idx))
                sentiment = event['sentiment_analysis']

                if sentiment.get('emotion_distribution'):
                    lines.append("**情感分布：**")
                    for emotion, count in sentiment['emotion_distribution'].items():
                        lines.append(f"- {emotion}: {count}次")

                if sentiment.get('public_demands'):
                    lines.append("\n**公众诉求：**")
                    for demand in sentiment['public_demands']:
                        lines.append(f"- {demand}")
                lines.append("")

            # 处置建议
            if event.get('recommendations'):
                lines.append("#### 3.{}.4 处置建议".format(idx))
                recs = event['recommendations']

                if recs.get('immediate'):
                    lines.append("**立即措施：**")
                    for rec in recs['immediate']:
                        lines.append(f"- {rec}")

                if recs.get('short_term'):
                    lines.append("\n**短期措施：**")
                    for rec in recs['short_term']:
                        lines.append(f"- {rec}")
                lines.append("")

            lines.append("---\n")

        return lines

    def _format_sentiment_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化情感分析部分"""
        lines = []
        sentiment = data.get('sentiment', {})

        lines.append("## 四、情感分析与公众关切\n")

        # 情感分布
        lines.append("### 4.1 情感倾向分析\n")
        emotion_dist = sentiment.get('emotion_distribution', {})

        if emotion_dist:
            lines.append("| 情感 | 占比 | 典型表述 |")
            lines.append("|------|------|----------|")
            total = sum(emotion_dist.values())
            for emotion, count in emotion_dist.items():
                percentage = (count / total * 100) if total > 0 else 0
                lines.append(f"| {emotion} | {percentage:.1f}% | 见关键词 |")
        else:
            lines.append("暂无详细情感分析数据\n")

        # 关键词
        lines.append("\n### 4.2 关键词云图\n")
        keywords = sentiment.get('top_keywords', [])
        wordcloud_image = sentiment.get('wordcloud_image')

        if wordcloud_image:
            lines.append(f"![关键词云图]({wordcloud_image})\n")

        if keywords:
            lines.append("**高频词（TOP 20）：**\n")
            lines.append("```")
            for kw in keywords[:20]:
                lines.append(f"{kw['keyword']} ({kw['count']}次)")
            lines.append("```")

        # 公众诉求
        lines.append("\n### 4.3 公众主要诉求\n")
        demands = sentiment.get('public_demands', [])

        if demands:
            for idx, demand in enumerate(demands, 1):
                lines.append(f"{idx}. {demand}")
        else:
            lines.append("暂无明确诉求")

        lines.append("\n---\n")

        return lines

    def _format_risk_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化风险评估部分"""
        lines = []
        risk = data.get('risk_assessment', {})

        lines.append("## 五、风险预警与评估\n")

        # 当前风险点
        lines.append("### 5.1 当前风险点\n")
        current_risks = risk.get('current_risks', [])

        if current_risks:
            lines.append("| 风险点 | 等级 | 平均风险分 | 事件数 |")
            lines.append("|--------|------|-----------|--------|")
            for r in current_risks[:5]:
                lines.append(f"| {r['department']} | {r['level_text']} | {r['avg_risk_score']} | {r['event_count']} |")
        else:
            lines.append("暂无风险点\n")

        # 风险等级统计
        lines.append("\n### 5.2 风险等级统计\n")
        risk_levels = risk.get('risk_levels', {})

        if risk_levels.get('red'):
            lines.append(f"**红色预警（极高）：** {', '.join(risk_levels['red'])}")
        if risk_levels.get('orange'):
            lines.append(f"**橙色预警（高）：** {', '.join(risk_levels['orange'])}")
        if risk_levels.get('yellow'):
            lines.append(f"**黄色预警（中）：** {', '.join(risk_levels['yellow'])}")

        lines.append("\n---\n")

        return lines

    def _format_recommendations_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化建议部分"""
        lines = []
        recs = data.get('recommendations', {})

        lines.append("## 六、应对措施与处置建议\n")

        # 立即措施
        immediate = recs.get('immediate_actions', [])
        if immediate:
            lines.append("### 6.1 立即应对措施（24小时内）\n")
            for action in immediate:
                lines.append(f"{action}")

        # 短期措施
        short_term = recs.get('short_term_actions', [])
        if short_term:
            lines.append("\n### 6.2 短期应对措施（1周内）\n")
            for action in short_term:
                lines.append(f"{action}")

        # 长期措施
        long_term = recs.get('long_term_actions', [])
        if long_term:
            lines.append("\n### 6.3 长期预防措施（1-3个月）\n")
            for action in long_term:
                lines.append(f"{action}")

        # 监测重点
        monitoring = recs.get('monitoring', {})
        if monitoring:
            lines.append("\n### 6.4 舆情监测重点\n")
            lines.append(f"**监测频率：** {monitoring.get('frequency', '')}")

            keywords = monitoring.get('keywords', [])
            if keywords:
                lines.append("\n**关键词：**")
                for kw in keywords[:10]:
                    lines.append(f"- {kw}")

        lines.append("\n---\n")

        return lines

    def _format_impact_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化影响预测部分"""
        lines = []
        impact = data.get('impact_forecast', {})

        lines.append("## 七、影响预测\n")

        # 短期
        short = impact.get('short_term', [])
        if short:
            lines.append("### 7.1 短期影响（1-7天）\n")
            for item in short:
                lines.append(f"- {item}")

        # 中期
        medium = impact.get('medium_term', [])
        if medium:
            lines.append("\n### 7.2 中期影响（1-4周）\n")
            for item in medium:
                lines.append(f"- {item}")

        # 长期
        long = impact.get('long_term', [])
        if long:
            lines.append("\n### 7.3 长期影响（1-3个月）\n")
            for item in long:
                lines.append(f"- {item}")

        # 法律风险
        legal = impact.get('legal_risk', {})
        if legal:
            lines.append("\n### 7.4 法律风险评估\n")
            lines.append(f"- **诉讼概率：** {legal.get('probability', '')}")
            lines.append(f"- **预估金额：** {legal.get('estimated_amount', '')}")
            lines.append(f"- **风险说明：** {legal.get('description', '')}")

        lines.append("\n---\n")

        return lines

    def _format_templates_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化声明模板部分"""
        lines = []
        templates = data.get('response_templates', {})

        lines.append("## 八、官方声明模板\n")

        # 首次回应
        first = templates.get('first_response', '')
        if first:
            lines.append("### 8.1 首次回应模板\n")
            lines.append("```")
            lines.append(first)
            lines.append("```")

        # 进展通报
        progress = templates.get('progress_update', '')
        if progress:
            lines.append("\n### 8.2 调查进展模板\n")
            lines.append("```")
            lines.append(progress)
            lines.append("```")

        lines.append("\n---\n")

        return lines

    def _format_appendix_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化附录部分"""
        lines = []
        appendix = data.get('appendix', {})

        lines.append("## 九、附录\n")

        # 事件清单
        event_list = appendix.get('event_list', [])
        if event_list:
            lines.append("### 附录A：负面舆情事件清单\n")
            lines.append("| ID | 时间 | 平台 | 类型 | 科室 | 风险分 | 状态 |")
            lines.append("|----|------|------|------|------|--------|------|")
            for event in event_list[:20]:
                lines.append(f"| {event['id']} | {event['time'][:10]} | {event['platform']} | {event['type']} | {event['department']} | {event['risk_score']} | {event['status']} |")

        # 联系方式
        contact = appendix.get('contact_info', {})
        if contact:
            lines.append("\n### 附录B：联系方式\n")
            monitoring = contact.get('monitoring_center', {})
            if monitoring:
                lines.append("**舆情监测中心：**")
                lines.append(f"- 负责人：{monitoring.get('负责人', '')}")
                lines.append(f"- 电话：{monitoring.get('电话', '')}")
                lines.append(f"- 邮箱：{monitoring.get('邮箱', '')}")

            crisis = contact.get('crisis_team', {})
            if crisis:
                lines.append("\n**危机公关小组：**")
                lines.append(f"- 组长：{crisis.get('组长', '')}")
                lines.append(f"- 成员：{', '.join(crisis.get('成员', []))}")

        lines.append("\n---\n")

        return lines

    def _generate_sentiment_analysis_new(self, df: pd.DataFrame) -> Dict[str, Any]:
        """
        新增功能1：情感分析与舆情态势
        - 正面/中性/负面舆情的比例分布
        - 情绪强度评分
        - 舆情热度趋势图数据
        """
        if len(df) == 0:
            return {
                'sentiment_distribution': {},
                'emotion_intensity': {},
                'trend_data': [],
                'sentiment_ratio': {}
            }

        # 1. 情感倾向分类（正面/中性/负面）
        sentiment_counts = {'positive': 0, 'neutral': 0, 'negative': 0}

        for _, row in df.iterrows():
            risk_score = row.get('风险分_数值', 0)
            severity = row.get('严重程度', 'low')

            # 根据风险分和严重程度判断情感倾向
            if risk_score >= 70 or severity == 'high':
                sentiment_counts['negative'] += 1
            elif risk_score >= 40 or severity == 'medium':
                sentiment_counts['neutral'] += 1
            else:
                sentiment_counts['positive'] += 1

        total = len(df)
        sentiment_ratio = {
            'negative': round(sentiment_counts['negative'] / total * 100, 1) if total > 0 else 0,
            'neutral': round(sentiment_counts['neutral'] / total * 100, 1) if total > 0 else 0,
            'positive': round(sentiment_counts['positive'] / total * 100, 1) if total > 0 else 0
        }

        # 2. 情绪强度评分（基于关键词和风险分）
        emotion_intensity = {}
        all_content = ' '.join(df.get('标题', pd.Series()).fillna('') + ' ' + df.get('内容', pd.Series()).fillna(''))

        for emotion, keywords in self.emotion_keywords.items():
            intensity = 0
            if JIEBA_AVAILABLE:
                words = list(jieba.cut(all_content))
                count = sum(1 for word in words if word in keywords)
                # 归一化到0-100
                intensity = min(100, count * 5)
            emotion_intensity[emotion] = intensity

        # 3. 舆情热度趋势图数据（按日期统计）
        daily_data = (
            df.groupby('日期', as_index=False)
            .agg(
                event_count=('风险分_数值', 'count'),
                avg_risk_score=('风险分_数值', 'mean'),
                max_risk_score=('风险分_数值', 'max'),
            )
        )

        trend_data = []
        for _, row in daily_data.iterrows():
            date = row['日期']
            count = int(row['event_count'])
            avg_risk = round(row['avg_risk_score'], 1)
            max_risk = int(row['max_risk_score'])

            # 确保日期是字符串格式
            if hasattr(date, 'strftime'):
                date_str = date.strftime('%Y-%m-%d')
            elif hasattr(date, 'isoformat'):
                date_str = date.isoformat()
            else:
                date_str = str(date)

            trend_data.append({
                'date': date_str,
                'count': count,
                'avg_risk_score': avg_risk,
                'max_risk_score': max_risk,
                'heat_index': count * avg_risk / 10  # 热度指数 = 数量 × 平均风险分 / 10
            })

        # 4. 公众情绪指数（综合指标）
        public_emotion_index = round(
            sentiment_ratio['negative'] * 1.0 +
            sentiment_ratio['neutral'] * 0.5 +
            sentiment_ratio['positive'] * 0.1,
            1
        )

        return {
            'sentiment_distribution': sentiment_counts,
            'sentiment_ratio': sentiment_ratio,
            'emotion_intensity': emotion_intensity,
            'trend_data': trend_data,
            'public_emotion_index': public_emotion_index,
            'emotion_level': self._assess_emotion_level(public_emotion_index)
        }

    def _assess_emotion_level(self, index: float) -> str:
        """评估情绪等级"""
        if index >= 80:
            return "极度负面"
        elif index >= 60:
            return "负面"
        elif index >= 40:
            return "偏负面"
        elif index >= 20:
            return "中性"
        else:
            return "正面"

    def _generate_category_statistics(self, df: pd.DataFrame) -> Dict[str, Any]:
        """
        新增功能3：舆情分类统计
        - 按问题类型分类统计（服务、收费、环境、医疗质量等）
        - 饼图/柱状图数据
        """
        if len(df) == 0:
            return {
                'categories': {},
                'chart_data': [],
                'top_categories': []
            }

        # 定义分类规则
        category_rules = {
            '医疗质量': ['死亡', '去世', '抢救', '手术', '误诊', '漏诊', '医疗事故', '并发症'],
            '服务态度': ['态度', '服务', '冷漠', '不耐烦', '推诿', '拒诊', '不会看病'],
            '收费问题': ['收费', '费用', '贵', '乱收费', '重复收费', '价格', '住院费'],
            '环境卫生': ['环境', '卫生', '脏', '垃圾', '厕所', '异味', '清洁'],
            '流程问题': ['等待', '排队', '时间长', '挂号', '预约', '流程'],
            '无障碍设施': ['无障碍', '视障', '残疾', '盲道', '轮椅', '导盲'],
            '隐私保护': ['隐私', '泄露', '病历', '信息', '个人信息'],
            '其他': []
        }

        # 分类统计
        category_counts = {cat: 0 for cat in category_rules.keys()}
        category_details = {cat: [] for cat in category_rules.keys()}

        for _, row in df.iterrows():
            content = str(row.get('标题', '')) + str(row.get('内容', '')) + str(row.get('警示理由', ''))
            risk_score = row.get('风险分_数值', 0)

            matched = False
            for category, keywords in category_rules.items():
                if category == '其他':
                    continue
                if any(keyword in content for keyword in keywords):
                    category_counts[category] += 1
                    category_details[category].append({
                        'title': row.get('标题', '')[:30],
                        'risk_score': int(risk_score),
                        'date': row.get('日期_字符串', '')
                    })
                    matched = True
                    break

            if not matched:
                category_counts['其他'] += 1
                category_details['其他'].append({
                    'title': row.get('标题', '')[:30],
                    'risk_score': int(risk_score),
                    'date': row.get('日期_字符串', '')
                })

        # 计算百分比和平均风险分
        total = len(df)
        categories = {}
        for cat, count in category_counts.items():
            if count > 0:
                details = category_details[cat]
                avg_risk = sum(d['risk_score'] for d in details) / len(details) if details else 0

                categories[cat] = {
                    'count': count,
                    'percentage': round(count / total * 100, 1),
                    'avg_risk_score': round(avg_risk, 1),
                    'severity': self._get_category_severity(avg_risk),
                    'details': details[:5]  # 只保留前5个示例
                }

        # 生成图表数据
        chart_data = []
        for cat, info in sorted(categories.items(), key=lambda x: x[1]['count'], reverse=True):
            chart_data.append({
                'category': cat,
                'count': info['count'],
                'percentage': info['percentage'],
                'color': self._get_category_color(cat)
            })

        # Top分类
        top_categories = sorted(
            [(cat, info) for cat, info in categories.items()],
            key=lambda x: x[1]['count'],
            reverse=True
        )[:5]

        return {
            'categories': categories,
            'chart_data': chart_data,
            'top_categories': [{'name': cat, **info} for cat, info in top_categories],
            'total_categories': len([c for c in categories.values() if c['count'] > 0])
        }

    def _get_category_severity(self, avg_risk: float) -> str:
        """获取分类严重程度"""
        if avg_risk >= 80:
            return '极高'
        elif avg_risk >= 60:
            return '高'
        elif avg_risk >= 40:
            return '中'
        else:
            return '低'

    def _get_category_color(self, category: str) -> str:
        """获取分类对应的颜色（用于图表）"""
        color_map = {
            '医疗质量': '#FF4444',
            '服务态度': '#FF8800',
            '收费问题': '#FFBB33',
            '环境卫生': '#00C851',
            '流程问题': '#33B5E5',
            '无障碍设施': '#AA66CC',
            '隐私保护': '#FF6F00',
            '其他': '#999999'
        }
        return color_map.get(category, '#999999')

    def _generate_spread_forecast(self, df: pd.DataFrame) -> Dict[str, Any]:
        """
        新增功能6：风险传播预测
        - 基于当前趋势的未来3-7天传播预测
        - 可能引发二次传播的风险点
        - 敏感时间节点提醒
        """
        if len(df) == 0:
            return {
                'forecast_3days': {},
                'forecast_7days': {},
                'secondary_risks': [],
                'sensitive_dates': []
            }

        # 1. 计算传播速度和趋势
        df_sorted = df.sort_values('创建时间_解析')
        if len(df_sorted) < 2:
            spread_rate = 0
        else:
            time_span = (df_sorted.iloc[-1]['创建时间_解析'] - df_sorted.iloc[0]['创建时间_解析']).total_seconds() / 86400  # 天数
            spread_rate = len(df) / max(time_span, 1)  # 每天新增数量

        # 2. 预测未来3天和7天
        current_count = len(df)
        avg_risk = df['风险分_数值'].mean()

        # 根据趋势调整预测系数
        trend = self._analyze_trend(df)
        if '上升' in trend:
            growth_factor = 1.5
        elif '下降' in trend:
            growth_factor = 0.7
        else:
            growth_factor = 1.0

        forecast_3days = {
            'predicted_new_events': int(spread_rate * 3 * growth_factor),
            'predicted_total_events': int(current_count + spread_rate * 3 * growth_factor),
            'predicted_avg_risk': round(avg_risk * (1.1 if '上升' in trend else 0.95), 1),
            'confidence': '中' if len(df) >= 5 else '低'
        }

        forecast_7days = {
            'predicted_new_events': int(spread_rate * 7 * growth_factor),
            'predicted_total_events': int(current_count + spread_rate * 7 * growth_factor),
            'predicted_avg_risk': round(avg_risk * (1.2 if '上升' in trend else 0.9), 1),
            'confidence': '中' if len(df) >= 5 else '低'
        }

        # 3. 识别二次传播风险点
        secondary_risks = []

        # 高风险事件可能引发二次传播
        high_risk_events = df[df['风险分_数值'] >= 80]
        if len(high_risk_events) > 0:
            secondary_risks.append({
                'risk_type': '高风险事件发酵',
                'description': f'存在{len(high_risk_events)}个高风险事件，可能引发媒体关注和二次传播',
                'probability': '高',
                'impact': '严重'
            })

        # 多平台传播风险
        platforms = df['来源_标准'].nunique()
        if platforms >= 3:
            secondary_risks.append({
                'risk_type': '跨平台传播',
                'description': f'舆情已在{platforms}个平台传播，存在跨平台发酵风险',
                'probability': '中',
                'impact': '较大'
            })

        # 集中爆发风险
        daily_counts = df.groupby('日期').size()
        if len(daily_counts) > 0 and daily_counts.max() >= len(df) * 0.5:
            secondary_risks.append({
                'risk_type': '集中爆发',
                'description': '舆情在短时间内集中爆发，可能持续发酵',
                'probability': '高',
                'impact': '严重'
            })

        # 情感极化风险
        negative_ratio = len(df[df['风险分_数值'] >= 70]) / len(df) if len(df) > 0 else 0
        if negative_ratio >= 0.7:
            secondary_risks.append({
                'risk_type': '情感极化',
                'description': f'{negative_ratio*100:.0f}%的舆情为高负面，公众情绪极化严重',
                'probability': '高',
                'impact': '严重'
            })

        # 4. 敏感时间节点提醒
        sensitive_dates = []
        now = datetime.now()

        # 周末（舆情容易发酵）
        days_to_weekend = (5 - now.weekday()) % 7
        if days_to_weekend <= 3:
            weekend_date = now + timedelta(days=days_to_weekend)
            sensitive_dates.append({
                'date': weekend_date.strftime('%Y-%m-%d'),
                'type': '周末',
                'reason': '周末期间用户活跃度高，舆情容易发酵传播',
                'suggestion': '加强周末舆情监控'
            })

        # 节假日（需要根据实际情况添加）
        # 这里简化处理，可以根据需要添加具体节假日判断

        # 月初/月末（医院就诊高峰）
        if now.day <= 5:
            sensitive_dates.append({
                'date': now.strftime('%Y-%m-%d'),
                'type': '月初',
                'reason': '月初就诊高峰期，医疗纠纷风险增加',
                'suggestion': '加强医患沟通，提升服务质量'
            })
        elif now.day >= 25:
            sensitive_dates.append({
                'date': now.strftime('%Y-%m-%d'),
                'type': '月末',
                'reason': '月末就诊高峰期，医疗纠纷风险增加',
                'suggestion': '加强医患沟通，提升服务质量'
            })

        # 5. 传播路径预测
        spread_path_prediction = self._predict_spread_path(df)

        return {
            'forecast_3days': forecast_3days,
            'forecast_7days': forecast_7days,
            'secondary_risks': secondary_risks,
            'sensitive_dates': sensitive_dates,
            'spread_path_prediction': spread_path_prediction,
            'overall_trend': trend,
            'spread_rate': round(spread_rate, 2)
        }

    def _predict_spread_path(self, df: pd.DataFrame) -> Dict[str, Any]:
        """预测传播路径"""
        if len(df) == 0:
            return {}

        # 分析当前传播路径
        platforms = df['来源_标准'].value_counts().to_dict()

        # 预测可能的传播路径
        predicted_platforms = []

        if '抖音' in platforms or '微博' in platforms:
            predicted_platforms.append({
                'platform': '传统媒体',
                'probability': '中',
                'reason': '社交媒体热度高，可能引起传统媒体关注'
            })

        if len(platforms) >= 2:
            predicted_platforms.append({
                'platform': '其他社交平台',
                'probability': '高',
                'reason': '已在多平台传播，可能扩散到更多平台'
            })

        return {
            'current_platforms': list(platforms.keys()),
            'predicted_platforms': predicted_platforms,
            'spread_pattern': '病毒式传播' if len(platforms) >= 3 else '线性传播'
        }

    def _format_sentiment_analysis_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化情感分析与舆情态势部分"""
        lines = []
        sentiment_new = data.get('sentiment_analysis_new', {})

        if not sentiment_new:
            return lines

        lines.append("## 四、情感分析与舆情态势\n")

        # 插入情感倾向饼图
        chart_paths = data.get('chart_paths', {})
        if chart_paths.get('sentiment_pie'):
            lines.append(f"![情感倾向分布]({chart_paths['sentiment_pie']})\n")

        # 1. 情感倾向分布
        lines.append("### 情感倾向分布\n")
        sentiment_ratio = sentiment_new.get('sentiment_ratio', {})
        sentiment_counts = sentiment_new.get('sentiment_distribution', {})

        if sentiment_ratio:
            lines.append("| 情感类型 | 数量 | 占比 | 说明 |")
            lines.append("|---------|------|------|------|")
            lines.append(f"| 负面 | {sentiment_counts.get('negative', 0)}条 | {sentiment_ratio.get('negative', 0)}% | 高风险舆情 |")
            lines.append(f"| 中性 | {sentiment_counts.get('neutral', 0)}条 | {sentiment_ratio.get('neutral', 0)}% | 中等风险舆情 |")
            lines.append(f"| 正面 | {sentiment_counts.get('positive', 0)}条 | {sentiment_ratio.get('positive', 0)}% | 低风险舆情 |")
            lines.append("")

        # 2. 公众情绪指数
        emotion_index = sentiment_new.get('public_emotion_index', 0)
        emotion_level = sentiment_new.get('emotion_level', '未知')
        # 去掉emoji
        emotion_level_clean = emotion_level.replace('🔴', '').replace('🟠', '').replace('🟡', '').replace('🟢', '').replace('🔵', '').strip()
        lines.append(f"**公众情绪指数：** {emotion_index}/100 - {emotion_level_clean}\n")

        # 3. 情绪强度分析
        emotion_intensity = sentiment_new.get('emotion_intensity', {})
        if emotion_intensity:
            lines.append("### 情绪强度分析\n")
            lines.append("| 情绪类型 | 强度值 | 等级 |")
            lines.append("|---------|--------|------|")
            for emotion, intensity in sorted(emotion_intensity.items(), key=lambda x: x[1], reverse=True):
                level = "高" if intensity >= 60 else "中" if intensity >= 30 else "低"
                lines.append(f"| {emotion} | {intensity}/100 | {level} |")
            lines.append("")

        # 4. 舆情热度趋势
        trend_data = sentiment_new.get('trend_data', [])
        if trend_data:
            lines.append("### 舆情热度趋势\n")

            # 插入趋势折线图
            chart_paths = data.get('chart_paths', {})
            if chart_paths.get('trend_line'):
                lines.append(f"![舆情热度趋势]({chart_paths['trend_line']})\n")
            lines.append("| 日期 | 舆情数量 | 平均风险分 | 热度指数 |")
            lines.append("|------|---------|-----------|---------|")
            for item in trend_data[-7:]:  # 只显示最近7天
                lines.append(f"| {item['date']} | {item['count']}条 | {item['avg_risk_score']} | {item['heat_index']:.1f} |")
            lines.append("")

        lines.append("---\n")
        return lines

    def _format_category_statistics_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化舆情分类统计部分"""
        lines = []
        category_stats = data.get('category_statistics', {})

        if not category_stats:
            return lines

        lines.append("## 五、舆情分类统计\n")

        # 插入问题类型饼图
        chart_paths = data.get('chart_paths', {})
        if chart_paths.get('category_pie'):
            lines.append(f"![问题类型分布]({chart_paths['category_pie']})\n")

        # 1. 分类概览
        categories = category_stats.get('categories', {})
        total_categories = category_stats.get('total_categories', 0)

        lines.append(f"**问题类型总数：** {total_categories}类\n")

        # 2. 分类详情表格
        lines.append("### 问题类型分布\n")
        if categories:
            lines.append("| 问题类型 | 数量 | 占比 | 平均风险分 | 严重程度 |")
            lines.append("|---------|------|------|-----------|---------|")

            for cat, info in sorted(categories.items(), key=lambda x: x[1]['count'], reverse=True):
                lines.append(f"| {cat} | {info['count']}条 | {info['percentage']}% | {info['avg_risk_score']} | {info['severity']} |")
            lines.append("")

        # 3. Top问题类型
        top_categories = category_stats.get('top_categories', [])
        if top_categories:
            lines.append("### 重点关注类型\n")
            for idx, cat_info in enumerate(top_categories[:3], 1):
                name = cat_info.get('name', '未知')
                count = cat_info.get('count', 0)
                avg_risk = cat_info.get('avg_risk_score', 0)
                details = cat_info.get('details', [])
                
                lines.append(f"**{idx}. {name}** ({count}条，平均风险分{avg_risk})")
                if details:
                    lines.append("   - 典型案例：")
                    for detail in details[:3]:
                        lines.append(f"     - {detail['title']} (风险分:{detail['risk_score']})")
                lines.append("")

        lines.append("---\n")
        return lines

    def _format_spread_forecast_section(self, data: Dict[str, Any]) -> List[str]:
        """格式化风险传播预测部分"""
        lines = []
        forecast = data.get('spread_forecast', {})

        if not forecast:
            return lines

        lines.append("## 八、风险传播预测\n")
        
        # 1. 传播趋势
        overall_trend = forecast.get('overall_trend', '未知')
        spread_rate = forecast.get('spread_rate', 0)
        lines.append(f"**当前传播趋势：** {overall_trend}")
        lines.append(f"**传播速度：** {spread_rate}条/天\n")
        
        # 2. 未来3天预测
        forecast_3days = forecast.get('forecast_3days', {})
        if forecast_3days:
            lines.append("### 8.1 未来3天预测\n")
            lines.append("| 指标 | 预测值 | 置信度 |")
            lines.append("|------|--------|--------|")
            lines.append(f"| 新增舆情数量 | {forecast_3days.get('predicted_new_events', 0)}条 | {forecast_3days.get('confidence', '低')} |")
            lines.append(f"| 累计舆情总数 | {forecast_3days.get('predicted_total_events', 0)}条 | {forecast_3days.get('confidence', '低')} |")
            lines.append(f"| 预测平均风险分 | {forecast_3days.get('predicted_avg_risk', 0)} | {forecast_3days.get('confidence', '低')} |")
            lines.append("")
        
        # 3. 未来7天预测
        forecast_7days = forecast.get('forecast_7days', {})
        if forecast_7days:
            lines.append("### 8.2 未来7天预测\n")
            lines.append("| 指标 | 预测值 | 置信度 |")
            lines.append("|------|--------|--------|")
            lines.append(f"| 新增舆情数量 | {forecast_7days.get('predicted_new_events', 0)}条 | {forecast_7days.get('confidence', '低')} |")
            lines.append(f"| 累计舆情总数 | {forecast_7days.get('predicted_total_events', 0)}条 | {forecast_7days.get('confidence', '低')} |")
            lines.append(f"| 预测平均风险分 | {forecast_7days.get('predicted_avg_risk', 0)} | {forecast_7days.get('confidence', '低')} |")
            lines.append("")
        
        # 4. 二次传播风险点
        secondary_risks = forecast.get('secondary_risks', [])
        if secondary_risks:
            lines.append("### 8.3 二次传播风险点\n")
            lines.append("| 风险类型 | 描述 | 发生概率 | 影响程度 |")
            lines.append("|---------|------|---------|---------|")
            for risk in secondary_risks:
                lines.append(f"| {risk['risk_type']} | {risk['description']} | {risk['probability']} | {risk['impact']} |")
            lines.append("")
        
        # 5. 敏感时间节点
        sensitive_dates = forecast.get('sensitive_dates', [])
        if sensitive_dates:
            lines.append("### 8.4 敏感时间节点提醒\n")
            for date_info in sensitive_dates:
                lines.append(f"- **{date_info['date']}** ({date_info['type']})")
                lines.append(f"  - 原因：{date_info['reason']}")
                lines.append(f"  - 建议：{date_info['suggestion']}")
            lines.append("")
        
        # 6. 传播路径预测
        spread_path = forecast.get('spread_path_prediction', {})
        if spread_path:
            lines.append("### 8.5 传播路径预测\n")
            current_platforms = spread_path.get('current_platforms', [])
            predicted_platforms = spread_path.get('predicted_platforms', [])
            spread_pattern = spread_path.get('spread_pattern', '未知')
            
            lines.append(f"**当前传播平台：** {', '.join(current_platforms)}")
            lines.append(f"**传播模式：** {spread_pattern}\n")
            
            if predicted_platforms:
                lines.append("**可能扩散平台：**")
                for platform in predicted_platforms:
                    lines.append(f"- {platform['platform']} (概率:{platform['probability']}) - {platform['reason']}")
                lines.append("")
        
        lines.append("---\n")
        return lines

    def generate_charts(self, report_data: Dict[str, Any], hospital_name: str) -> Dict[str, str]:
        """
        生成所有图表并返回图片路径
        
        返回：
        {
            'sentiment_pie': '情感倾向饼图路径',
            'category_pie': '问题类型饼图路径',
            'trend_line': '舆情热度趋势图路径',
            'platform_pie': '平台分布饼图路径'
        }
        """
        if not MATPLOTLIB_AVAILABLE:
            return {}
        
        chart_paths = {}
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # 1. 情感倾向饼图
        sentiment_path = self._generate_sentiment_pie_chart(
            report_data, 
            os.path.join(self.charts_dir, f'sentiment_pie_{timestamp}.png')
        )
        if sentiment_path:
            chart_paths['sentiment_pie'] = sentiment_path
        
        # 2. 问题类型饼图
        category_path = self._generate_category_pie_chart(
            report_data,
            os.path.join(self.charts_dir, f'category_pie_{timestamp}.png')
        )
        if category_path:
            chart_paths['category_pie'] = category_path
        
        # 3. 舆情热度趋势图
        trend_path = self._generate_trend_line_chart(
            report_data,
            os.path.join(self.charts_dir, f'trend_line_{timestamp}.png')
        )
        if trend_path:
            chart_paths['trend_line'] = trend_path
        
        # 4. 平台分布饼图
        platform_path = self._generate_platform_pie_chart(
            report_data,
            os.path.join(self.charts_dir, f'platform_pie_{timestamp}.png')
        )
        if platform_path:
            chart_paths['platform_pie'] = platform_path
        
        return chart_paths

    def _generate_sentiment_pie_chart(self, report_data: Dict[str, Any], output_path: str) -> str:
        """生成情感倾向饼图"""
        try:
            sentiment_data = report_data.get('sentiment_analysis_new', {})
            sentiment_counts = sentiment_data.get('sentiment_distribution', {})
            
            if not sentiment_counts:
                return None
            
            # 准备数据
            labels = []
            sizes = []
            colors = []
            
            if sentiment_counts.get('negative', 0) > 0:
                labels.append('负面')
                sizes.append(sentiment_counts['negative'])
                colors.append('#FF4444')
            
            if sentiment_counts.get('neutral', 0) > 0:
                labels.append('中性')
                sizes.append(sentiment_counts['neutral'])
                colors.append('#FFBB33')
            
            if sentiment_counts.get('positive', 0) > 0:
                labels.append('正面')
                sizes.append(sentiment_counts['positive'])
                colors.append('#00C851')
            
            if not sizes:
                return None
            
            # 创建图表
            fig, ax = plt.subplots(figsize=(8, 6))
            wedges, texts, autotexts = ax.pie(
                sizes, 
                labels=labels, 
                colors=colors,
                autopct='%1.1f%%',
                startangle=90,
                textprops={'fontsize': 12}
            )
            
            # 设置标题
            ax.set_title('情感倾向分布', fontsize=16, fontweight='bold', pad=20)
            
            # 美化百分比文字
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(11)
            
            plt.tight_layout()
            plt.savefig(output_path, dpi=150, bbox_inches='tight')
            plt.close()
            
            return output_path
        except Exception as e:
            print(f"生成情感倾向饼图失败: {e}")
            return None

    def _generate_category_pie_chart(self, report_data: Dict[str, Any], output_path: str) -> str:
        """生成问题类型饼图"""
        try:
            category_data = report_data.get('category_statistics', {})
            categories = category_data.get('categories', {})
            
            if not categories:
                return None
            
            # 准备数据（只显示前6个类型）
            sorted_categories = sorted(categories.items(), key=lambda x: x[1]['count'], reverse=True)[:6]
            
            labels = []
            sizes = []
            colors = []
            
            color_map = {
                '医疗质量': '#FF4444',
                '服务态度': '#FF8800',
                '收费问题': '#FFBB33',
                '环境卫生': '#00C851',
                '流程问题': '#33B5E5',
                '无障碍设施': '#AA66CC',
                '隐私保护': '#FF6F00',
                '其他': '#999999'
            }
            
            for cat_name, cat_info in sorted_categories:
                labels.append(f"{cat_name}\n({cat_info['count']}条)")
                sizes.append(cat_info['count'])
                colors.append(color_map.get(cat_name, '#999999'))
            
            # 创建图表
            fig, ax = plt.subplots(figsize=(10, 7))
            wedges, texts, autotexts = ax.pie(
                sizes,
                labels=labels,
                colors=colors,
                autopct='%1.1f%%',
                startangle=90,
                textprops={'fontsize': 10}
            )
            
            # 设置标题
            ax.set_title('问题类型分布', fontsize=16, fontweight='bold', pad=20)
            
            # 美化百分比文字
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(10)
            
            plt.tight_layout()
            plt.savefig(output_path, dpi=150, bbox_inches='tight')
            plt.close()
            
            return output_path
        except Exception as e:
            print(f"生成问题类型饼图失败: {e}")
            return None

    def _generate_trend_line_chart(self, report_data: Dict[str, Any], output_path: str) -> str:
        """生成舆情热度趋势折线图"""
        try:
            sentiment_data = report_data.get('sentiment_analysis_new', {})
            trend_data = sentiment_data.get('trend_data', [])
            
            if not trend_data or len(trend_data) < 2:
                return None
            
            # 准备数据
            dates = [item['date'] for item in trend_data]
            counts = [item['count'] for item in trend_data]
            avg_risks = [item['avg_risk_score'] for item in trend_data]
            
            # 创建图表
            fig, ax1 = plt.subplots(figsize=(12, 6))
            
            # 左Y轴：舆情数量
            color1 = '#2196F3'
            ax1.set_xlabel('日期', fontsize=12)
            ax1.set_ylabel('舆情数量（条）', color=color1, fontsize=12)
            line1 = ax1.plot(dates, counts, color=color1, marker='o', linewidth=2, 
                            markersize=6, label='舆情数量')
            ax1.tick_params(axis='y', labelcolor=color1)
            ax1.grid(True, alpha=0.3, linestyle='--')
            
            # 右Y轴：平均风险分
            ax2 = ax1.twinx()
            color2 = '#FF5722'
            ax2.set_ylabel('平均风险分', color=color2, fontsize=12)
            line2 = ax2.plot(dates, avg_risks, color=color2, marker='s', linewidth=2,
                            markersize=6, linestyle='--', label='平均风险分')
            ax2.tick_params(axis='y', labelcolor=color2)
            
            # 设置标题
            plt.title('舆情热度趋势', fontsize=16, fontweight='bold', pad=20)
            
            # 旋转X轴标签
            plt.xticks(rotation=45, ha='right')
            
            # 添加图例
            lines = line1 + line2
            labels = [l.get_label() for l in lines]
            ax1.legend(lines, labels, loc='upper left', fontsize=10)
            
            plt.tight_layout()
            plt.savefig(output_path, dpi=150, bbox_inches='tight')
            plt.close()
            
            return output_path
        except Exception as e:
            print(f"生成舆情热度趋势图失败: {e}")
            return None

    def _generate_platform_pie_chart(self, report_data: Dict[str, Any], output_path: str) -> str:
        """生成平台分布饼图"""
        try:
            overview = report_data.get('overview', {})
            platform_dist = overview.get('platform_distribution', {})
            
            if not platform_dist:
                return None
            
            # 准备数据
            labels = []
            sizes = []
            colors = ['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0', '#9966FF', '#FF9F40']
            
            for idx, (platform, count) in enumerate(sorted(platform_dist.items(), 
                                                          key=lambda x: x[1], 
                                                          reverse=True)):
                labels.append(f"{platform}\n({count}条)")
                sizes.append(count)
            
            # 创建图表
            fig, ax = plt.subplots(figsize=(9, 7))
            wedges, texts, autotexts = ax.pie(
                sizes,
                labels=labels,
                colors=colors[:len(sizes)],
                autopct='%1.1f%%',
                startangle=90,
                textprops={'fontsize': 11}
            )
            
            # 设置标题
            ax.set_title('平台分布', fontsize=16, fontweight='bold', pad=20)
            
            # 美化百分比文字
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(11)
            
            plt.tight_layout()
            plt.savefig(output_path, dpi=150, bbox_inches='tight')
            plt.close()
            
            return output_path
        except Exception as e:
            print(f"生成平台分布饼图失败: {e}")
            return None
