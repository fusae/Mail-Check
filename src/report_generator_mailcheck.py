#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
舆情报告生成器 - Mail-Check专用版本
从MySQL数据库读取舆情数据，自动生成详细分析报告
"""

import db
import pandas as pd
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Any, Optional
import os
import json
import tempfile
import subprocess

try:
    from wordcloud import WordCloud
    WORDCLOUD_AVAILABLE = True
except ImportError:
    WORDCLOUD_AVAILABLE = False


_ENHANCED_GENERATOR_CLS = None


def _load_enhanced_generator_cls():
    """延迟加载增强版报告生成器，避免每次请求重复修改sys.path。"""
    global _ENHANCED_GENERATOR_CLS
    if _ENHANCED_GENERATOR_CLS is not None:
        return _ENHANCED_GENERATOR_CLS

    import sys
    current_dir = os.path.dirname(os.path.abspath(__file__))
    if current_dir not in sys.path:
        sys.path.insert(0, current_dir)

    from report_generator_enhanced import EnhancedReportGenerator
    _ENHANCED_GENERATOR_CLS = EnhancedReportGenerator
    return _ENHANCED_GENERATOR_CLS


class MailCheckReportGenerator:
    """Mail-Check舆情报告生成器（MySQL版）"""

    def __init__(self, db_path: str = None):
        """
        初始化报告生成器

        Args:
            db_path: 已废弃（保留兼容，不再使用）
        """
        current_dir = os.path.dirname(os.path.abspath(__file__))
        self.project_root = os.path.dirname(current_dir)
        self.db_path = db_path
        self.conn = None

    def connect(self):
        """连接数据库"""
        if not self.conn:
            self.conn = db.connect(self.project_root)
        return self.conn

    def close(self):
        """关闭数据库连接"""
        if self.conn:
            self.conn.close()
            self.conn = None

    def query(self, sql: str, params: tuple = ()):
        """执行查询"""
        conn = self.connect()
        cursor = self.conn.cursor()
        cursor.execute(sql, params)
        rows = cursor.fetchall()
        return rows

    def _normalize_hospital_filter(self, hospital: Optional[str]) -> Optional[str]:
        if not isinstance(hospital, str):
            return hospital
        cleaned = hospital.strip()
        if cleaned in ("all", "全部", "全院汇总", "all hospitals"):
            return None
        return cleaned or None

    def _format_report_date_range(
        self,
        start_date: Optional[str],
        end_date: Optional[str]
    ) -> Optional[str]:
        """优先保留用户选择的时间范围，用于报告封面展示。"""
        def _format_date(value: Optional[str]) -> Optional[str]:
            if not value:
                return None
            try:
                return datetime.strptime(value[:10], "%Y-%m-%d").strftime("%Y年%m月%d日")
            except ValueError:
                return value

        start_text = _format_date(start_date)
        end_text = _format_date(end_date)

        if start_text and end_text:
            return f"{start_text}-{end_text}"
        if start_text:
            return f"{start_text}起"
        if end_text:
            return f"截至{end_text}"
        return None

    def fetch_all_data(
        self,
        start_date: str = None,
        end_date: str = None,
        hospital: str = None,
        include_dismissed: bool = False,
        dedupe_by_event: bool = True,
        sentiment_ids: Optional[List[str]] = None,
        record_ids: Optional[List[int]] = None
    ):
        """
        获取所有负面舆情数据

        Args:
            start_date: 开始日期 (YYYY-MM-DD)
            end_date: 结束日期 (YYYY-MM-DD)
            hospital: 医院名称（可选，筛选特定医院）
            include_dismissed: 是否包含已标记误报的数据
            dedupe_by_event: 是否按事件归并去重（同事件仅保留最高风险且最新一条）
            sentiment_ids: 指定纳入报告的舆情ID列表（可选）
            record_ids: 指定纳入报告的数据库记录ID列表（可选，优先精确筛选）

        Returns:
            DataFrame: 舆情数据
        """
        # 清洗筛选参数
        hospital = self._normalize_hospital_filter(hospital)

        if isinstance(start_date, str) and len(start_date) >= 10:
            start_date = start_date[:10]
        if isinstance(end_date, str) and len(end_date) >= 10:
            end_date = end_date[:10]

        # 构建SQL查询
        sql = """
            SELECT
                sentiment_id as ID,
                id as 记录ID,
                hospital_name as 医院,
                title as 标题,
                source as 来源,
                severity as 严重程度,
                processed_at as 创建时间,
                reason as 警示理由,
                content as 内容,
                url as 原文链接,
                status as 状态,
                event_id as 事件ID,
                is_duplicate as 重复事件
            FROM negative_sentiments
            WHERE 1=1
        """
        params = []

        # 日期筛选
        if start_date:
            sql += " AND DATE(processed_at) >= ?"
            params.append(start_date)
        if end_date:
            sql += " AND DATE(processed_at) <= ?"
            params.append(end_date)

        # 医院筛选
        if hospital:
            sql += " AND hospital_name = ?"
            params.append(hospital)

        if record_ids is not None:
            clean_record_ids = []
            for rid in record_ids:
                try:
                    clean_record_ids.append(int(rid))
                except Exception:
                    continue
            if not clean_record_ids:
                return pd.DataFrame()
            placeholders = ",".join(["?"] * len(clean_record_ids))
            sql += f" AND id IN ({placeholders})"
            params.extend(clean_record_ids)

        # 勾选导出：只纳入用户选中的舆情
        if sentiment_ids is not None:
            clean_ids = [str(sid).strip() for sid in sentiment_ids if str(sid).strip()]
            if not clean_ids:
                return pd.DataFrame()
            placeholders = ",".join(["?"] * len(clean_ids))
            sql += f" AND sentiment_id IN ({placeholders})"
            params.extend(clean_ids)

        # 默认排除已误报项，让报告聚焦“待处置”舆情
        if not include_dismissed:
            sql += " AND COALESCE(NULLIF(status, ''), 'active') != 'dismissed'"

        sql += " ORDER BY processed_at DESC"

        # 执行查询
        rows = self.query(sql, tuple(params))

        # 转换为DataFrame
        data = []
        for row in rows:
            data.append({
                'ID': row['ID'],
                '记录ID': row.get('记录ID'),
                '医院': row['医院'],
                '标题': row['标题'],
                '来源': row['来源'],
                '严重程度': row['严重程度'],
                '创建时间': row['创建时间'],
                '警示理由': row['警示理由'],
                '内容': row['内容'],
                '原文链接': row['原文链接'],
                '状态': row['状态'] or 'active',
                '事件ID': row.get('事件ID'),
                '重复事件': int(row.get('重复事件') or 0)
            })

        df = pd.DataFrame(data)

        if df.empty:
            return df

        # 添加风险分
        if '严重程度' in df.columns:
            df['风险分'] = df['严重程度'].apply(
                lambda x: 100 if x == 'high' else 60 if x == 'medium' else 30
            )

        if dedupe_by_event:
            df = self._dedupe_event_rows(df)

        return df

    def _dedupe_event_rows(self, df: pd.DataFrame) -> pd.DataFrame:
        """按事件归并去重：同事件保留最高风险、最新创建时间的一条。"""
        if df.empty:
            return df

        df = df.copy()
        df['创建时间_解析'] = pd.to_datetime(df['创建时间'], errors='coerce')
        df['事件键'] = df['事件ID'].apply(
            lambda x: f"event:{int(x)}" if pd.notna(x) else None
        )
        df['事件键'] = df['事件键'].fillna(df['ID'].apply(lambda x: f"item:{x}"))

        # 排序优先级：风险分高 -> 时间新
        df = df.sort_values(
            by=['风险分', '创建时间_解析'],
            ascending=[False, False],
            na_position='last'
        )
        deduped = df.drop_duplicates(subset=['事件键'], keep='first').copy()
        deduped = deduped.sort_values(by='创建时间_解析', ascending=False, na_position='last')

        return deduped.drop(columns=['创建时间_解析', '事件键'], errors='ignore')

    def generate_report(
        self,
        start_date: str = None,
        end_date: str = None,
        hospital: str = None,
        report_period: str = None,
        output_format: str = 'markdown',
        include_dismissed: bool = False,
        dedupe_by_event: bool = True,
        sentiment_ids: Optional[List[str]] = None,
        record_ids: Optional[List[int]] = None
    ):
        """
        生成舆情报告

        Args:
            start_date: 开始日期
            end_date: 结束日期
            hospital: 医院名称
            report_period: 报告周期（如"2026年第一季度"）
            output_format: 输出格式（markdown/word）
            include_dismissed: 是否包含已误报数据
            dedupe_by_event: 是否按事件归并去重
            sentiment_ids: 仅纳入指定舆情ID
            record_ids: 仅纳入指定数据库记录ID

        Returns:
            dict: 包含报告数据和文件路径
        """
        print(f"[INFO] 正在获取数据...")
        hospital_filter = self._normalize_hospital_filter(hospital)
        raw_df = self.fetch_all_data(
            start_date=start_date,
            end_date=end_date,
            hospital=hospital_filter,
            include_dismissed=include_dismissed,
            dedupe_by_event=False,
            sentiment_ids=sentiment_ids,
            record_ids=record_ids
        )
        if len(raw_df) == 0:
            print("[WARN] 没有找到符合条件的数据")
            return {
                'success': False,
                'message': '没有数据'
            }
        df = self._dedupe_event_rows(raw_df) if dedupe_by_event else raw_df

        print(f"[INFO] 获取到 {len(raw_df)} 条舆情数据，报告纳入 {len(df)} 条")

        # 导入增强版report_generator
        try:
            enhanced_cls = _load_enhanced_generator_cls()
            gen = enhanced_cls()
        except ImportError:
            print("[ERROR] 无法导入EnhancedReportGenerator，请确保report_generator_enhanced.py在src目录下")
            return {
                'success': False,
                'message': '缺少report_generator_enhanced模块'
            }

        output_format = (output_format or 'markdown').strip().lower()
        if output_format in ('md',):
            output_format = 'markdown'
        elif output_format in ('docx',):
            output_format = 'word'
        elif output_format not in ('markdown', 'word', 'both'):
            output_format = 'markdown'

        # 确定医院名称
        if hospital_filter:
            hospital_name = hospital_filter
        elif len(df) > 0:
            # 如果只有一个医院，使用它；否则使用通用名称
            unique_hospitals = df['医院'].unique()
            if len(unique_hospitals) == 1:
                hospital_name = unique_hospitals[0]
            else:
                hospital_name = "多医院汇总"
        else:
            hospital_name = "未知"

        # 确定报告周期
        if not report_period:
            if start_date and end_date:
                report_period = f"{start_date} 至 {end_date}"
            else:
                report_period = datetime.now().strftime('%Y年%m月')
        report_date_range = self._format_report_date_range(start_date, end_date)

        # 生成报告数据
        print(f"[INFO] 正在分析数据...")
        report_data = gen.generate_report_data(
            df=df,
            hospital_name=hospital_name,
            report_type='special',
            report_period=report_period,
            report_date_range=report_date_range
        )
        report_data['data_scope'] = {
            'include_dismissed': include_dismissed,
            'dedupe_by_event': dedupe_by_event,
            'raw_count': int(len(raw_df)),
            'included_count': int(len(df)),
            'excluded_count': int(len(raw_df) - len(df))
        }

        # 生成报告文件
        output_dir = Path(os.path.join(self.project_root, 'data', 'reports'))
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_hospital_name = ''.join(ch if ch.isalnum() or ch in ('-', '_') else '_' for ch in hospital_name)
        filename = f"{safe_hospital_name}_舆情报告_{timestamp}"

        # 生成关键词云（可选）
        wordcloud_name = self._generate_wordcloud(report_data, output_dir, filename)
        if wordcloud_name:
            report_data.setdefault('sentiment', {})['wordcloud_image'] = wordcloud_name
        wordcloud_path = output_dir / wordcloud_name if wordcloud_name else None

        result = {
            'success': True,
            'hospital_name': hospital_name,
            'period': report_period,
            'total_events': len(df),
            'raw_total_events': len(raw_df),
            'high_risk_events': len(df[df['严重程度'] == 'high']),
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'data_scope': report_data.get('data_scope', {}),
            'included_hospitals': sorted([str(x) for x in df['医院'].dropna().unique().tolist()]) if '医院' in df.columns else [],
            'files': {}
        }

        # 先生成Markdown内容（Word将基于Markdown转换）
        md_content = gen.generate_markdown_report(report_data)

        # Markdown格式
        if output_format in ['markdown', 'md', 'both']:
            md_path = output_dir / f"{filename}.md"
            print(f"[INFO] 正在生成Markdown报告...")
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            result['files']['markdown'] = str(md_path)
            print(f"[OK] Markdown报告: {md_path}")

        # Word格式（基于Markdown转换）
        if output_format in ['word', 'docx', 'both']:
            docx_path = output_dir / f"{filename}.docx"
            print(f"[INFO] 正在生成Word报告...")

            # 如果没有保存Markdown文件，则使用临时文件
            temp_md_path = None
            if 'markdown' in result['files']:
                md_for_docx = result['files']['markdown']
            else:
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.md', dir=str(output_dir))
                temp_file.write(md_content.encode('utf-8'))
                temp_file.flush()
                temp_file.close()
                temp_md_path = temp_file.name
                md_for_docx = temp_md_path

            self._convert_markdown_to_docx(md_for_docx, str(docx_path))
            if wordcloud_path and wordcloud_path.exists():
                self._embed_wordcloud_in_docx(str(docx_path), str(wordcloud_path))
            result['files']['word'] = str(docx_path)
            print(f"[OK] Word报告: {docx_path}")

            if temp_md_path:
                try:
                    os.remove(temp_md_path)
                except OSError:
                    pass

        # 添加摘要
        result['summary'] = report_data['summary']
        result['overview'] = report_data['overview']

        self.close()
        return result

    def _find_chinese_font(self) -> Optional[str]:
        """尝试寻找系统中文字体路径"""
        candidates = [
            # macOS
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
            # Linux (common)
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/arphic/ukai.ttc",
            "/usr/share/fonts/truetype/arphic/uming.ttc",
            # Windows (if mounted)
            "C:\\Windows\\Fonts\\msyh.ttc",
            "C:\\Windows\\Fonts\\simhei.ttf",
        ]
        for path in candidates:
            if os.path.exists(path):
                return path
        return None

    def _generate_wordcloud(self, report_data: Dict[str, Any], output_dir: Path, filename: str) -> Optional[str]:
        """根据关键词生成词云图片，返回相对文件名"""
        if not WORDCLOUD_AVAILABLE:
            print("[WARN] 未安装wordcloud，跳过关键词云生成")
            return None

        sentiment = report_data.get('sentiment', {})
        keywords = sentiment.get('top_keywords', [])
        if not keywords:
            return None

        font_path = self._find_chinese_font()
        if not font_path:
            print("[WARN] 未找到中文字体，跳过关键词云生成")
            return None

        freqs: Dict[str, int] = {}
        for item in keywords:
            if isinstance(item, dict):
                key = item.get('keyword')
                count = item.get('count', 1)
            else:
                try:
                    key, count = item
                except Exception:
                    continue
            if not key:
                continue
            key_text = str(key).replace('\n', ' ').replace('\r', ' ')
            key_text = ' '.join(key_text.split()).strip()
            if not key_text:
                continue
            freqs[key_text] = int(count) if count else 1

        if not freqs:
            return None

        wc = WordCloud(
            font_path=font_path,
            width=1200,
            height=800,
            background_color="white",
            colormap="viridis",
            prefer_horizontal=1.0
        )
        wc.generate_from_frequencies(freqs)

        image_name = f"{filename}_wordcloud.png"
        image_path = output_dir / image_name
        wc.to_file(str(image_path))
        return image_name

    def _convert_markdown_to_docx(self, md_path: str, docx_path: str) -> None:
        """将Markdown转换为Word（优先pypandoc，其次调用pandoc命令）"""
        # 优先使用pypandoc
        try:
            import pypandoc  # type: ignore

            pypandoc.convert_file(md_path, 'docx', outputfile=docx_path)
            return
        except Exception:
            pass

        # 退回使用pandoc命令行
        try:
            subprocess.run(['pandoc', md_path, '-o', docx_path], check=True)
            return
        except FileNotFoundError as exc:
            raise RuntimeError("未找到pandoc，请先安装pandoc或pypandoc。") from exc
        except subprocess.CalledProcessError as exc:
            raise RuntimeError(f"pandoc转换失败: {exc}") from exc

    def _embed_wordcloud_in_docx(self, docx_path: str, image_path: str) -> None:
        """将关键词云图片插入到Word文档中"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.text.paragraph import Paragraph
            from docx.oxml import OxmlElement
        except ImportError:
            print("[WARN] 未安装python-docx，无法嵌入关键词云图片")
            return

        def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            return new_para

        doc = Document(docx_path)
        target = None
        for p in doc.paragraphs:
            if "关键词云图" in p.text:
                target = p
                break

        if target:
            pic_par = _insert_paragraph_after(target)
            pic_par.add_run().add_picture(image_path, width=Inches(5.5))
        else:
            doc.add_heading("关键词云图", level=3)
            doc.add_picture(image_path, width=Inches(5.5))

        doc.save(docx_path)


def main():
    """命令行入口"""
    import argparse

    parser = argparse.ArgumentParser(description='Mail-Check舆情报告生成器')
    parser.add_argument('--start-date', help='开始日期 (YYYY-MM-DD)')
    parser.add_argument('--end-date', help='结束日期 (YYYY-MM-DD)')
    parser.add_argument('--hospital', help='医院名称')
    parser.add_argument('--period', help='报告周期（如"2026年第一季度"）')
    parser.add_argument('--format', default='markdown', choices=['markdown', 'word', 'both'], help='输出格式')
    parser.add_argument('--db', help='数据库路径（可选）')
    parser.add_argument('--include-dismissed', action='store_true', help='包含已标记误报数据')
    parser.add_argument('--no-dedupe', action='store_true', help='不按事件归并去重')

    args = parser.parse_args()

    # 创建报告生成器
    generator = MailCheckReportGenerator(db_path=args.db)

    # 生成报告
    print("="*60)
    print("  Mail-Check 舆情报告生成器")
    print("="*60)
    print()

    result = generator.generate_report(
        start_date=args.start_date,
        end_date=args.end_date,
        hospital=args.hospital,
        report_period=args.period,
        output_format=args.format,
        include_dismissed=args.include_dismissed,
        dedupe_by_event=not args.no_dedupe
    )

    print()
    print("="*60)
    if result['success']:
        print("  报告生成成功！")
        print("="*60)
        print(f"医院: {result.get('hospital_name', 'N/A')}")
        print(f"周期: {result.get('period', 'N/A')}")
        print(f"事件总数: {result.get('total_events', 0)}")
        print(f"高风险事件: {result.get('high_risk_events', 0)}")
        print()
        print("生成的文件:")
        for fmt, path in result.get('files', {}).items():
            print(f"  [{fmt}] {path}")
    else:
        print("  报告生成失败")
        print("="*60)
        print(f"原因: {result.get('message', '未知错误')}")

    print()
    print("="*60)


if __name__ == "__main__":
    main()
